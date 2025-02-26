import pandas as pd
import docx
from docx import Document
from docx.shared import Inches, Pt, Twips, RGBColor
from docx.oxml import OxmlElement, parse_xml
from docx.oxml.ns import qn, nsdecls
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_ALIGN_VERTICAL
import requests
from io import BytesIO
from PIL import Image, UnidentifiedImageError  # Added proper PIL import
import sys
import os
import logging
import urllib.parse
import tempfile  # Added for temporary file handling

# Enhanced logging setup
logging.basicConfig(
    level=logging.DEBUG,
    format='%(asctime)s - %(name)s - %(levelname)s - %(message)s',
    handlers=[
        logging.StreamHandler(),
        logging.FileHandler('patent_formatter.log')
    ]
)
logger = logging.getLogger(__name__)

class PatentDocumentFormatter:
    # Constants remain the same...
    LEFT_COLUMN_WIDTH = Inches(1.38)
    RIGHT_COLUMN_WIDTH = Inches(5.61)
    STANDARD_ROW_HEIGHT = Inches(0.24)
    HEADING_FONT_SIZE = Pt(11)
    STANDARD_FONT_SIZE = Pt(10)
    PAGE_WIDTH = Twips(12240)
    PAGE_HEIGHT = Twips(15840)
    MARGIN = Twips(1440)
    
    def __init__(self, template_path):
        try:
            if not os.path.exists(template_path):
                raise FileNotFoundError(f"Template file not found: {template_path}")
            self.doc = Document(template_path)
            self._setup_document_properties()
            logger.info("Document formatter initialized successfully")
        except Exception as e:
            logger.error(f"Error initializing document formatter: {e}")
            raise

    def _setup_document_properties(self):
        """Setup document properties based on XML specifications"""
        section = self.doc.sections[0]
        # Set page margins
        section.top_margin = self.MARGIN
        section.bottom_margin = self.MARGIN
        section.left_margin = self.MARGIN
        section.right_margin = self.MARGIN
        # Set page size
        section.page_width = self.PAGE_WIDTH
        section.page_height = self.PAGE_HEIGHT
        
        # Set default font
        style = self.doc.styles['Normal']
        style.font.name = 'Calibri'
        style.font.size = self.STANDARD_FONT_SIZE
        
        # Setup ListParagraph style
        if 'ListParagraph' not in self.doc.styles:
            list_style = self.doc.styles.add_style('ListParagraph', docx.enum.style.WD_STYLE_TYPE.PARAGRAPH)
            list_style.base_style = self.doc.styles['Normal']
            list_style.font.name = 'Calibri'
            list_style.font.size = self.STANDARD_FONT_SIZE

    def set_cell_border(self, cell):
        """Add borders to cell"""
        tc = cell._tc
        tcPr = tc.get_or_add_tcPr()
        
        borders = parse_xml(f'''
            <w:tcBorders xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main">
                <w:top w:val="single" w:sz="4" w:space="0" w:color="auto"/>
                <w:left w:val="single" w:sz="4" w:space="0" w:color="auto"/>
                <w:bottom w:val="single" w:sz="4" w:space="0" w:color="auto"/>
                <w:right w:val="single" w:sz="4" w:space="0" w:color="auto"/>
            </w:tcBorders>
        ''')
        
        existing_borders = tcPr.find(qn('w:tcBorders'))
        if existing_borders is not None:
            tcPr.remove(existing_borders)
        
        tcPr.append(borders)

    def set_cell_width(self, cell, width):
        """Set exact cell width"""
        tc = cell._tc
        tcPr = tc.get_or_add_tcPr()
        tcW = OxmlElement('w:tcW')
        tcW.set(qn('w:w'), str(width.twips))
        tcW.set(qn('w:type'), 'dxa')
        
        for element in tcPr.findall(qn('w:tcW')):
            tcPr.remove(element)
        
        tcPr.append(tcW)

    def set_row_height(self, row, height):
        """Set exact row height"""
        tr = row._tr
        trPr = tr.get_or_add_trPr()
        trHeight = OxmlElement('w:trHeight')
        trHeight.set(qn('w:val'), str(height.twips))
        trHeight.set(qn('w:hRule'), 'exact')
        
        for element in trPr.findall(qn('w:trHeight')):
            trPr.remove(element)
            
        trPr.append(trHeight)

    def add_hyperlink(self, paragraph, url):
        """Add hyperlink with Calibri font"""
        if not url or pd.isna(url):
            run = paragraph.add_run("No Link Available")
            run.font.name = 'Calibri'
            run.font.size = self.STANDARD_FONT_SIZE
            return

        part = paragraph.part
        r_id = part.relate_to(url, docx.opc.constants.RELATIONSHIP_TYPE.HYPERLINK, is_external=True)
        
        run = paragraph.add_run("Link")
        run.font.name = 'Calibri'
        run.font.size = self.STANDARD_FONT_SIZE
        run.font.underline = True
        run.font.color.rgb = RGBColor(0, 0, 255)
        
        hyperlink = OxmlElement('w:hyperlink')
        hyperlink.set(qn('r:id'), r_id)
        
        hyperlink.append(run._element)
        paragraph._p.append(hyperlink)

    def process_image(self, image_url, max_width):
        """
        Download, resize, and return an image for insertion in Word.
        Uses a temporary file to avoid Windows file locking issues.
        """
        if not image_url or pd.isna(image_url):
            logger.warning("No image URL provided")
            return None

        try:
            # Clean and encode the URL properly
            parsed_url = urllib.parse.urlparse(image_url)
            cleaned_url = urllib.parse.urlunparse(
                parsed_url._replace(
                    path=urllib.parse.quote(parsed_url.path),
                    query=urllib.parse.quote(parsed_url.query, safe='=&')
                )
            )

            logger.debug(f"Fetching image from: {cleaned_url}")

            headers = {
                'User-Agent': 'Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36',
                'Accept': 'image/jpeg,image/png,image/*'
            }

            # Use a named temporary file
            with tempfile.NamedTemporaryFile(delete=False, suffix='.jpg') as tmp_file:
                temp_path = tmp_file.name  # Get the temporary file path
                
                response = requests.get(cleaned_url, timeout=30, headers=headers, stream=True)
                response.raise_for_status()

                for chunk in response.iter_content(chunk_size=8192):
                    tmp_file.write(chunk)
            
            # Open the image and close it immediately after processing
            with Image.open(temp_path) as img:
                if img.mode in ('RGBA', 'LA'):
                    background = Image.new('RGB', img.size, 'white')
                    background.paste(img, mask=img.split()[-1])
                    img = background

                aspect_ratio = img.width / img.height
                new_width = min(max_width.inches * 72, img.width)
                new_height = new_width / aspect_ratio

                img = img.resize((int(new_width), int(new_height)), Image.LANCZOS)

                # Save the processed image back to the temporary file
                img.save(temp_path, format='JPEG', quality=95, optimize=True)

            return temp_path  # Return the temporary file path

        except requests.exceptions.SSLError as e:
            logger.error(f"SSL Error while fetching image: {e}")
        except requests.exceptions.RequestException as e:
            logger.error(f"Network error while fetching image: {e}")
        except UnidentifiedImageError as e:
            logger.error(f"Invalid image format: {e}")
        except Exception as e:
            logger.error(f"Unexpected error processing image: {e}")
        return None

    def add_section_heading(self, text):
        """Add section heading with Calibri font"""
        heading = self.doc.add_paragraph(style='ListParagraph')
        heading.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        run = heading.add_run(text)
        run.font.name = 'Calibri'
        run.font.size = self.HEADING_FONT_SIZE
        run.font.bold = True
        run.font.underline = True
        
        index_para = self.doc.add_paragraph(style='ListParagraph')
        index_para.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
        index_run = index_para.add_run("<< INDEX")
        index_run.font.name = 'Calibri'
        index_run.font.size = self.STANDARD_FONT_SIZE
        index_run.font.color.rgb = RGBColor(0, 0, 255)
        index_run.font.underline = True

    def format_table_cell(self, cell, field=None):
        """
        Apply consistent cell formatting with Calibri font
        Now handles special cases for Abstract and Image
        """
        tc = cell._tc
        tcPr = tc.get_or_add_tcPr()
        
        # Set vertical alignment based on field type
        if field in ["Abstract", "Image"]:
            # Top alignment for Abstract and Image
            tcPr.append(parse_xml(r'<w:vAlign %s w:val="top"/>' % nsdecls('w')))
        else:
            # Center alignment for other fields
            tcPr.append(parse_xml(r'<w:vAlign %s w:val="center"/>' % nsdecls('w')))
        
        # Add cell margins
        margins_xml = parse_xml(
            r'<w:tcMar %s>'
            r'<w:top w:w="120" w:type="dxa"/>'
            r'<w:bottom w:w="120" w:type="dxa"/>'
            r'<w:left w:w="120" w:type="dxa"/>'
            r'<w:right w:w="120" w:type="dxa"/>'
            r'</w:tcMar>' % nsdecls('w')
        )
        
        # Remove existing margins if any
        existing_margins = tcPr.find(qn('w:tcMar'))
        if existing_margins is not None:
            tcPr.remove(existing_margins)
        
        tcPr.append(margins_xml)
        self.set_cell_border(cell)
        
        for paragraph in cell.paragraphs:
            paragraph.style = self.doc.styles['ListParagraph']
            paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
            for run in paragraph.runs:
                run.font.name = 'Calibri'
                run.font.size = self.STANDARD_FONT_SIZE

    def create_patent_table(self, data_row, family_to_image, remaining_space):
        """Creates a patent table ensuring proper spacing and avoiding page breaks."""
        try:
            fields = [
                "Serial No", "Publication No", "Kind Code", "Title", 
                "Publication Date", "Earliest Priority Date", "Assignee", 
                "Inventors", "Category", "IPC", "PDF Document", "Abstract", "Image"
            ]
            
            table = self.doc.add_table(rows=len(fields), cols=2)
            table.allow_autofit = False
            
            family_number = data_row.get("Family number")
            image_url = family_to_image.get(family_number, None)  # Get image URL using family number
            
            for i, field in enumerate(fields):
                try:
                    if field == "Image":
                        logger.debug(f"Processing image for row with URL: {image_url}")
                        
                        img_path = self.process_image(image_url, self.RIGHT_COLUMN_WIDTH)
                        if img_path:
                            try:
                                # Check available space before inserting the image
                                self.check_and_add_page_break(self.doc, remaining_space, 500)  # Example height

                                # Resize the image only if necessary
                                img_path = self.resize_image_to_fit(img_path, self.RIGHT_COLUMN_WIDTH.twips, 500)

                                # Insert centered image
                                self.add_centered_image_with_spacing(table.cell(i, 1), img_path)

                                os.unlink(img_path)  # Delete temp file after inserting
                                logger.debug("Successfully added image and deleted temp file")
                            except Exception as e:
                                logger.error(f"Error adding image to document: {e}")
                                table.cell(i, 1).text = "Error Loading Image"
                        else:
                            table.cell(i, 1).text = "Image Not Available"
                    
                    # Handle other fields normally
                    else:
                        table.cell(i, 1).text = str(data_row.get(field, ""))
                        
                except Exception as e:
                    logger.error(f"Error processing field {field}: {e}")
                    raise

            # Prevent table row splitting
            self.prevent_row_splitting(table)

        except Exception as e:
            logger.error(f"Error creating patent table: {e}")
            raise

    def create_document(self, excel_path, output_path):
        """Create the complete patent document"""
        try:
            if not os.path.exists(excel_path):
                raise FileNotFoundError(f"Excel file not found: {excel_path}")
                
            logger.info(f"Starting document creation from {excel_path}")
            xl = pd.ExcelFile(excel_path)
            
            # Verify required sheets exist
            required_sheets = ["First Publication", "Granted", "Sheet1"]
            missing_sheets = [sheet for sheet in required_sheets if sheet not in xl.sheet_names]
            if missing_sheets:
                raise ValueError(f"Missing required sheets: {', '.join(missing_sheets)}")
            
            first_pub_df = xl.parse("First Publication")
            granted_df = xl.parse("Granted")
            images_df = xl.parse("Sheet1")
            
            # Create Family Number to Image URL Mapping
            family_to_image = dict(zip(images_df["Family number"], images_df["Image"]))
            
            self.add_section_heading("FIRST PUBLICATIONS")
            for idx, row in first_pub_df.iterrows():
                logger.debug(f"Processing First Publication row {idx}")
                self.create_patent_table(row, family_to_image, self.PAGE_HEIGHT.twips - self.MARGIN.twips)
                self.doc.add_page_break()
            
            self.add_section_heading("GRANTED PATENTS")
            for idx, row in granted_df.iterrows():
                logger.debug(f"Processing Granted Patent row {idx}")
                self.create_patent_table(row, family_to_image, self.PAGE_HEIGHT.twips - self.MARGIN.twips)
                self.doc.add_page_break()
            
            # Ensure output directory exists
            os.makedirs(os.path.dirname(os.path.abspath(output_path)), exist_ok=True)
            
            self.doc.save(output_path)
            logger.info(f"Document successfully created: {output_path}")
            
        except Exception as e:
            logger.error(f"Error creating document: {e}")
            raise

def main():
    if len(sys.argv) != 4:
        print("Usage: script.py <excel_path> <output_path> <template_path>")
        sys.exit(1)
        
    try:
        excel_path = sys.argv[1]
        output_path = sys.argv[2]
        template_path = sys.argv[3]
        
        formatter = PatentDocumentFormatter(template_path)
        formatter.create_document(excel_path, output_path)
        
    except Exception as e:
        print(f"Error: {str(e)}")
        sys.exit(1)

if __name__ == "__main__":
    main()