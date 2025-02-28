import pandas as pd
from docx import Document
from docx.shared import Pt, Inches
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.enum.table import WD_ALIGN_VERTICAL
import sys
import logging
import requests
from io import BytesIO
from PIL import Image

# Set up logging
logging.basicConfig(level=logging.INFO, format='%(asctime)s - %(levelname)s - %(message)s')
logger = logging.getLogger(__name__)

def add_hyperlink(paragraph, text, url):
    """Add a hyperlink to a paragraph."""
    part = paragraph.part
    r_id = part.relate_to(url, 'http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink', is_external=True)
    
    hyperlink = OxmlElement('w:hyperlink')
    hyperlink.set(qn('r:id'), r_id)
    
    run = OxmlElement('w:r')
    rPr = OxmlElement('w:rPr')
    
    # Add underline
    underline = OxmlElement('w:u')
    underline.set(qn('w:val'), 'single')
    rPr.append(underline)
    
    # Add blue color
    color = OxmlElement('w:color')
    color.set(qn('w:val'), '0000FF')
    rPr.append(color)
    
    run.append(rPr)
    t = OxmlElement('w:t')
    t.text = text
    run.append(t)
    hyperlink.append(run)
    paragraph._element.append(hyperlink)

def download_and_insert_image(cell, image_url):
    """Downloads an image from a URL and inserts it into a cell."""
    try:
        response = requests.get(image_url, stream=True)
        if response.status_code == 200:
            img = Image.open(BytesIO(response.content))
            temp_img = BytesIO()
            img.save(temp_img, format='PNG')
            temp_img.seek(0)
            img_para = cell.add_paragraph()
            img_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = img_para.add_run()
            run.add_picture(temp_img, width=Inches(2))
    except Exception as e:
        logger.error(f"Error downloading image: {e}")

def add_section_header(document, title):
    """Add a section header to the document."""
    heading_para = document.add_paragraph()
    heading_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    heading_run = heading_para.add_run(title)
    heading_run.bold = True
    heading_run.underline = True
    heading_run.font.size = Pt(11)
    
    # Add index
    index_para = document.add_paragraph()
    index_para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    index_run = index_para.add_run("<<INDEX")
    index_run.font.size = Pt(10)

def create_patent_table(document, record, headers, df_images):
    """Create a table for a patent record."""
    # Create a table with 2 columns
    table = document.add_table(rows=0, cols=2)
    table.style = 'Table Grid'
    
    # Set column widths
    table.columns[0].width = Inches(1.38)
    table.columns[1].width = Inches(5.61)
    
    # Add rows for each field
    for header in headers:
        row_cells = table.add_row().cells
        row_cells[0].text = header
        row_cells[0].vertical_alignment = WD_ALIGN_VERTICAL.TOP
        row_cells[1].text = str(record.get(header, "")) if pd.notna(record.get(header, "")) else ""
        row_cells[1].vertical_alignment = WD_ALIGN_VERTICAL.TOP
        
        # Set fixed height for all fields except Abstract and Image
        if header not in ["Abstract", "Image"]:
            row_cells[0].height = Inches(0.28)  # Updated to 0.28 as requested
            row_cells[1].height = Inches(0.28)  # Updated to 0.28 as requested
    
    # Add hyperlink for PDF Document/Patent Link
    link_field = 'Patent Link'
    if link_field in headers and pd.notna(record.get(link_field)):
        pdf_para = table.rows[headers.index(link_field)].cells[1].paragraphs[0]
        pdf_para.clear()
        add_hyperlink(pdf_para, "Link", record[link_field])
    
    # Lookup and add image if available
    image_url = None
    if "Family number" in record and pd.notna(record["Family number"]):
        image_row = df_images[df_images['Family number'] == record['Family number']]
        if not image_row.empty and 'Image' in image_row.iloc[0] and pd.notna(image_row.iloc[0]['Image']):
            image_url = image_row.iloc[0]['Image']

    if image_url:
        image_cells = table.add_row().cells
        image_cells[0].text = "Image"
        image_cells[0].vertical_alignment = WD_ALIGN_VERTICAL.TOP
        download_and_insert_image(image_cells[1], image_url)
    
    return table

def create_first_publications_section(document, df, df_images):
    """Create the First Publications section in the document."""
    # Add section header
    add_section_header(document, "FIRST PUBLICATIONS")
    
    # Define headers for First Publications
    headers = [
        "Serial No", "Publication No", "Kind Code", "Title", 
        "Publication Date", "Earliest Priority Date", "Assignee", 
        "Inventors", "Category", "IPC", "Patent Link", "Abstract"
    ]
    
    # Iterate over each record in the DataFrame
    first_record = True
    for _, row in df.iterrows():
        # Only add page break after the first record
        if not first_record:
            document.add_page_break()
            
            # Add index at the top of each new page
            index_para = document.add_paragraph()
            index_para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            index_run = index_para.add_run("<<INDEX")
            index_run.font.size = Pt(10)
        
        first_record = False
        
        # Create table for this record
        create_patent_table(document, row.to_dict(), headers, df_images)

def create_granted_patents_section(document, df_granted, df_images):
    """Create the Granted Patents section in the document."""
    # Insert page break before granted patents section
    document.add_page_break()
    
    # Add section header
    add_section_header(document, "GRANTED PATENTS")
    
    # Define headers for Granted Patents
    headers = [
        "Serial No", "Family number", "Patent No", "Kind Code",
        "Title", "Publication Date", "Earliest Priority", "Assignee",
        "Inventors", "Category", "IPC", "Patent Link", "Abstract"
    ]
    
    # Process each granted patent record
    first_record = True
    for _, row in df_granted.iterrows():
        # Only add page break after the first record
        if not first_record:
            document.add_page_break()
            
            # Add index at the top of each new page
            index_para = document.add_paragraph()
            index_para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            index_run = index_para.add_run("<<INDEX")
            index_run.font.size = Pt(10)
        
        first_record = False
        
        # Create table for this record
        create_patent_table(document, row.to_dict(), headers, df_images)

def main():
    try:
        # Load Excel sheets
        excel_path = sys.argv[1] if len(sys.argv) > 1 else r'C:\Users\Ayman\Documents\Abhijit_mail_attachments\Test_PW.xlsm'
        df_fp = pd.read_excel(excel_path, sheet_name="First Publication")
        df_granted = pd.read_excel(excel_path, sheet_name="Granted")
        df_images = pd.read_excel(excel_path, sheet_name="Sheet1")
        
        # Load the template document
        template_file = "basic_page_template.docx"
        document = Document(template_file)
        
        # Create the First Publications section
        create_first_publications_section(document, df_fp, df_images)
        
        # Create the Granted Patents section
        create_granted_patents_section(document, df_granted, df_images)
        
        # Save the final document
        output_path = "part_4.docx"
        document.save(output_path)
        logger.info(f"Document saved: {output_path}")
        print(f"Successfully created document: {output_path}")
    
    except Exception as e:
        logger.error(f"Error generating document: {e}")
        print(f"Critical error: {e}")

if __name__ == "__main__":
    main()