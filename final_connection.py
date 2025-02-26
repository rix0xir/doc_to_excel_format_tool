import pandas as pd
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.oxml import OxmlElement, parse_xml
from docx.oxml.ns import qn, nsdecls
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_UNDERLINE, WD_PARAGRAPH_ALIGNMENT
from docx.enum.table import WD_ALIGN_VERTICAL, WD_CELL_VERTICAL_ALIGNMENT
from datetime import datetime
from io import BytesIO
import requests
import math

# -----------------------------
# Helper Functions (for tables and formatting)
# -----------------------------

# Function to set table borders
def set_table_borders(table):
    tbl = table._tbl
    tblPr = tbl.tblPr
    borders = OxmlElement('w:tblBorders')
    for border_name in ['top', 'left', 'bottom', 'right', 'insideH', 'insideV']:
        border = OxmlElement(f'w:{border_name}')
        border.set(qn('w:val'), 'single')
        border.set(qn('w:sz'), '4')  # Border thickness
        border.set(qn('w:space'), '0')
        border.set(qn('w:color'), '000000')  # Black color for the border
        borders.append(border)
    tblPr.append(borders)

# Function to set cell size and alignment
def set_cell_size_and_alignment(table, set_width=True):
    for row in table.rows:
        for cell in row.cells:
            if set_width:
                cell.width = Inches(1.38)  # Set width if specified
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            for paragraph in cell.paragraphs:
                paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    table.alignment = WD_ALIGN_PARAGRAPH.RIGHT



# Function to add detailed table
def add_detailed_table(document, row_data, headings):
    table = document.add_table(rows=len(headings), cols=2)
    for i, heading in enumerate(headings):
        table.cell(i, 0).text = heading
        table.cell(i, 1).text = str(row_data.get(heading, ''))
    set_table_borders(table)

# Function to add hyperlinks to publication numbers
def add_hyperlinked_value(cell, publication_no):
    paragraph = cell.paragraphs[0]  # Ensure we're working with a paragraph

    # Create hyperlink pointing to the bookmark
    hyperlink = OxmlElement('w:hyperlink')
    hyperlink.set(qn('w:anchor'), f'pub_{publication_no}')  # Link to the bookmark

    # Display publication number as hyperlink
    run = paragraph.add_run(publication_no)
    run.font.size = Pt(10)
    run.font.color.rgb = RGBColor(0, 0, 255)
    run.underline = True

    # Append hyperlink to paragraph
    hyperlink.append(run._r)
    paragraph._p.append(hyperlink)




# Adjust cell height for company names and categories
def adjust_cell_heights(table, company_names, category_names):
    for row in table.rows:
        tr = row._tr
        trPr = tr.get_or_add_trPr()
        trHeight = OxmlElement('w:trHeight')
        for cell in row.cells:
            text = cell.text.strip().upper()
            if text in company_names:
                trHeight.set(qn('w:val'), str(int(0.37 * 1440)))  # 0.37 inches in twips
                trHeight.set(qn('w:hRule'), 'exact')
                trPr.append(trHeight)
                break  # Set row height once for the entire row
            elif text in category_names:
                trHeight.set(qn('w:val'), str(int(0.28 * 1440)))  # 0.28 inches in twips
                trHeight.set(qn('w:hRule'), 'exact')
                trPr.append(trHeight)
                break  # Set row height once for the entire row

def set_row_height(row, height_in_inches):
    tr = row._tr
    trPr = tr.get_or_add_trPr()
    trHeight = OxmlElement('w:trHeight')
    trHeight.set(qn('w:val'), str(int(height_in_inches * 1440)))  # Convert inches to twips
    trHeight.set(qn('w:hRule'), 'exact')  # Set exact height
    trPr.append(trHeight)


def set_cell_width(cell, width_twips):
    """
    Set the width of a table cell in twips.
    """
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    
    # Remove any existing width settings
    tcW = tcPr.find(qn('w:tcW'))
    if tcW is not None:
        tcPr.remove(tcW)
    
    # Set new width
    new_tcW = OxmlElement('w:tcW')
    new_tcW.set(qn('w:w'), str(width_twips))
    new_tcW.set(qn('w:type'), 'dxa')
    tcPr.append(new_tcW)

def set_cell_margins(cell, top=100, start=100, bottom=100, end=100):
    """
    Set the margins inside a table cell.
    Margins are specified in twips (1/20 of a point, 1440 twips = 1 inch).
    Default margins are set to 100 twips (~0.07 inch).
    """
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()

    # Remove existing margin settings if any
    for margin_tag in ('w:topMar', 'w:startMar', 'w:bottomMar', 'w:endMar'):
        existing_margin = tcPr.find(qn(margin_tag))
        if existing_margin is not None:
            tcPr.remove(existing_margin)

    # Create new margin settings
    cell_margins = OxmlElement('w:tcMar')
    cell_margins.set(qn('w:top'), str(top))
    cell_margins.set(qn('w:start'), str(start))
    cell_margins.set(qn('w:bottom'), str(bottom))
    cell_margins.set(qn('w:end'), str(end))

    # Append margins to cell properties
    tcPr.append(cell_margins)

def set_paragraph_font(paragraph, font_name='Calibri', font_size=Pt(10), bold=False, italic=False, underline=False):
    """
    Set the font style, size, and other properties for a paragraph.
    Default font is Calibri, size 10pt.
    """
    for run in paragraph.runs:
        run.font.name = font_name
        run.font.size = font_size
        run.bold = bold
        run.italic = italic
        run.underline = underline


# Helper to format dates
def format_date(date_str):
    try:
        date_obj = datetime.strptime(date_str, '%Y-%m-%d %H:%M:%S')
        return date_obj.strftime('%d/%b/%Y')
    except ValueError:
        return date_str  # Return original if not in expected format
# -----------------------------
# Part 2: Adding the First Two Pages (Patent Watch)
# -----------------------------

def add_first_two_pages(document, excel_path):
    # Add the main title
    title_paragraph = document.add_paragraph('2445_2446 - PATENT WATCH – (04-NOV-2024 to 15-NOV-2024)')
    title_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_run = title_paragraph.runs[0]
    title_run.bold = True
    title_run.font.size = Pt(14)
    title_run.font.all_caps = True

    document.add_paragraph()  # Line break

    # Add INDEX heading
    index_paragraph = document.add_paragraph('INDEX')
    index_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    index_run = index_paragraph.runs[0]
    index_run.bold = True
    index_run.font.size = Pt(12)
    index_run.font.all_caps = True
    index_run.underline = WD_UNDERLINE.SINGLE

    # Create the combined company table (5 on top, 6 below)
    first_row_data = ['CGG', 'WESTERNGECO', 'PGS', 'ION', 'BGP/CNPC/PETROCHINA']
    second_row_data = ['CHEVRON', 'HALLIBURTON/LANDMARK', 'FFA GEOTERIC', 'PARADIGM', 'WEATHERFORD', 'BAKER HUGHES']
    company_table = create_company_table(document, first_row_data, second_row_data)

    # Set row height to 0.37" for this specific table
    for row in company_table.rows:
        set_row_height_exact(row, 0.37)

    document.add_paragraph()  # Line break

    # Add First Publications and Granted Patents hyperlinks
    link_paragraph = document.add_paragraph()
    link_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

    first_pub_run = link_paragraph.add_run('First Publications')
    first_pub_run.font.size = Pt(10)
    first_pub_run.underline = True
    first_pub_run.font.color.rgb = RGBColor(0, 0, 255)

    link_paragraph.add_run('\t\t')  # Adding 2 tab spaces

    granted_pat_run = link_paragraph.add_run('Granted Patents')
    granted_pat_run.font.size = Pt(10)
    granted_pat_run.underline = True
    granted_pat_run.font.color.rgb = RGBColor(0, 0, 255)

    document.add_paragraph()  # Line break

    # Load the FP and Grant worksheets
    df_fp = pd.read_excel(excel_path, sheet_name='FP')
    df_fp.columns = df_fp.columns.str.strip()
    df_fp['Category'] = df_fp['Category'].astype(str).str.strip()

    df_grant = pd.read_excel(excel_path, sheet_name='Grant')
    df_grant.columns = df_grant.columns.str.strip()
    df_grant['Category'] = df_grant['Category'].astype(str).str.strip()

    # Define categories
    categories = ['Seafloor', 'Land', 'Marine', 'Microseismic & Multiphysics',
                  'Processing', 'Reservoir', 'Geology', 'Data Management & Computing']

    # Split categories into two groups: first five and remaining three
    categories_group1 = categories[:5]
    categories_group2 = categories[5:]

    # Process the first group of categories
    for category in categories_group1:
        add_category_with_values(document, category, df_fp, df_grant)

    # Insert the company table again
    document.add_paragraph()  # Line break
    company_table_2 = create_company_table(document, first_row_data, second_row_data)

    # Set row height to 0.28" for all tables after the first company table
    for row in company_table_2.rows:
        set_row_height_exact(row, 0.28)

    document.add_paragraph()  # Line break

    # Process the remaining categories
    for category in categories_group2:
        add_category_with_values(document, category, df_fp, df_grant)


def set_row_height_exact(row, height_in_inches):
    tr = row._tr
    trPr = tr.get_or_add_trPr()
    trHeight = OxmlElement('w:trHeight')
    trHeight.set(qn('w:val'), str(int(height_in_inches * 1440)))  # Convert inches to twips
    trHeight.set(qn('w:hRule'), 'exact')  # Set exact height
    trPr.append(trHeight)

# Helper function to create the company table
def create_company_table(document, first_row_data, second_row_data):
    table = document.add_table(rows=2, cols=6)

    # First row
    for i, company in enumerate(first_row_data):
        cell = table.rows[0].cells[i]
        cell.text = company
        cell.paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        cell.paragraphs[0].runs[0].font.size = Pt(10)

    # Merge extra cell if fewer than 6 companies
    if len(first_row_data) < 6:
        table.rows[0].cells[-2].merge(table.rows[0].cells[-1])

    # Second row
    for i, company in enumerate(second_row_data):
        cell = table.rows[1].cells[i]
        cell.text = company
        cell.paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        cell.paragraphs[0].runs[0].font.size = Pt(10)

    set_table_borders(table)
    table.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    table.autofit = True

    return table  # Ensure the function returns the table object







# Helper function to add categories with FP and Grant values
def add_category_with_values(document, category, df_fp, df_grant):
    # Add bullet point for the category
    bullet_paragraph = document.add_paragraph()
    bullet_paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
    bullet_run = bullet_paragraph.add_run(u'\u2192 ' + category)  # Unicode for arrow
    bullet_run.font.size = Pt(10)
    bullet_run.bold = True

    # Extract Publication No from FP and Patent No from Grant for the category
    fp_values = df_fp[df_fp['Category'].str.lower().str.contains(category.lower())]['Publication No'].tolist()
    grant_values = df_grant[df_grant['Category'].str.lower().str.contains(category.lower())]['Patent No'].tolist()

    values = fp_values + grant_values

    # Create a 4-column table to display the values
    if values:
        num_rows = math.ceil(len(values) / 4)
        value_table = document.add_table(rows=num_rows, cols=4)
        value_table.autofit = True

        for idx, value in enumerate(values):
            row_idx = idx // 4
            col_idx = idx % 4
            cell = value_table.rows[row_idx].cells[col_idx]
            add_hyperlinked_value(cell, str(value))  # Add hyperlinks to publication numbers
            cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

        set_table_borders(value_table)
        set_cell_size_and_alignment(value_table)
        adjust_cell_heights(value_table, [], [category.upper()])  # Adjust heights for category names
    else:
        no_value_paragraph = document.add_paragraph(f'No records found for {category}')
        no_value_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

    document.add_paragraph()  # Line break after each category

# -----------------------------
# Part 3: Adding the First Publications Section
# -----------------------------

def add_first_publications_section(document, excel_path):
    # Load data from the 'FP' worksheet
    df_fp = pd.read_excel(excel_path, sheet_name='FP')
    df_fp.columns = df_fp.columns.str.strip()
    df_fp['Category'] = df_fp['Category'].astype(str).str.strip()

    # Define categories
    categories = [
        'Seafloor', 'Land', 'Marine', 'Microseismic & Multiphysics',
        'Processing', 'Reservoir', 'Geology', 'Data Management & Computing', 'Downhole'
    ]

    # Add main heading
    heading_paragraph = document.add_paragraph('FIRST PUBLICATIONS')
    heading_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    heading_run = heading_paragraph.runs[0]
    heading_run.bold = True
    heading_run.font.size = Pt(12)

    document.add_paragraph()  # Line break

    # Create index table with categories
    index_table = document.add_table(rows=1, cols=len(categories))
    index_row = index_table.rows[0].cells

    for i, category in enumerate(categories):
        paragraph = index_row[i].paragraphs[0]
        run = paragraph.add_run(category)
        run.font.size = Pt(10)
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

    set_table_borders(index_table)
    # Set alignment to center
    index_table.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Set cell height to 0.22 inches for the index table
    for row in index_table.rows:
        for cell in row.cells:
            tc = cell._tc
            tcPr = tc.get_or_add_tcPr()
            tcHeight = OxmlElement('w:trHeight')
            tcHeight.set(qn('w:val'), str(int(0.22 * 1440)))  # 0.22 inches in twips
            tcHeight.set(qn('w:hRule'), 'exact')
            tcPr.append(tcHeight)


    document.add_paragraph()  # Line break

    # Create the main data table
    create_fp_data_table(document, df_fp, categories)


# Helper function to create the FP data table
def create_fp_data_table(document, df_fp, categories):
    # Create main data table with 5 columns
    table = document.add_table(rows=1, cols=5)
    table.autofit = False  # Disable autofit to set custom widths

    # Set column widths
    column_widths = [Inches(0.4), Inches(1.08), Inches(2.46), Inches(1.38), Inches(1.48)]

    # Set column headers
    headers = ['Sl No', 'Publication No', 'Title', 'Assignee', 'Inventors']
    header_row = table.rows[0].cells

    for i, (header, width) in enumerate(zip(headers, column_widths)):
        cell = header_row[i]
        cell.text = header
        cell.width = width
        paragraph = cell.paragraphs[0]
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        paragraph.runs[0].bold = True
        paragraph.runs[0].font.size = Pt(10)

    # Process each category and add corresponding data
    for category in categories:
        # Add category heading row
        cat_row = table.add_row()
        cat_cell = cat_row.cells[0]
        cat_cell.merge(cat_row.cells[-1])
        cat_cell.text = f'> {category.upper()}'
        paragraph = cat_cell.paragraphs[0]
        paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
        paragraph.runs[0].bold = True
        paragraph.runs[0].font.size = Pt(10)

        # Filter data for the category
        cat_data = df_fp[df_fp['Category'].str.lower().str.contains(category.lower())]

        # Add rows for the data
        for _, row in cat_data.iterrows():
            data_row = table.add_row()
            data_row.cells[0].text = str(row.get('Serial No', ''))
            data_row.cells[1].text = str(row.get('Publication No', ''))
            data_row.cells[2].text = str(row.get('Title', ''))
            data_row.cells[3].text = str(row.get('Assignee', ''))
            data_row.cells[4].text = str(row.get('Inventors', ''))

            # Set column widths and font size for data rows
            for i, cell in enumerate(data_row.cells):
                cell.width = column_widths[i]
                for paragraph in cell.paragraphs:
                    paragraph.space_after = Pt(0)
                    paragraph.space_before = Pt(0)
                    for run in paragraph.runs:
                        run.font.size = Pt(10)

    set_table_borders(table)
# -----------------------------
# Part 4: Adding the Granted Patents Section
# -----------------------------

def add_granted_patents_section(document, excel_path):
    # Load data from the 'Grant' worksheet
    df_grant = pd.read_excel(excel_path, sheet_name='Grant')
    df_grant.columns = df_grant.columns.str.strip()
    df_grant['Category'] = df_grant['Category'].astype(str).str.strip()

    # Define categories
    categories = [
        'Seafloor', 'Land', 'Marine', 'Microseismic & Multiphysics',
        'Processing', 'Reservoir', 'Geology', 'Data Management & Computing', 'Downhole'
    ]

    # Add main heading
    heading_paragraph = document.add_paragraph('GRANTED PATENTS')
    heading_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    heading_run = heading_paragraph.runs[0]
    heading_run.bold = True
    heading_run.font.size = Pt(12)

    document.add_paragraph()  # Line break

    # Create index table with categories
    index_table = document.add_table(rows=1, cols=len(categories))
    index_row = index_table.rows[0].cells

    for i, category in enumerate(categories):
        cell = index_row[i]
        
        # Set cell width to allow more space for text (adjust width as needed)
        cell.width = Inches(1.5)  
        
        # Enable text wrapping by ensuring text stays inside the cell
        tc = cell._tc
        tcPr = tc.get_or_add_tcPr()
        tcPr.append(OxmlElement('w:noWrap'))  # Disable noWrap to allow text wrapping

        # Add the category text
        paragraph = cell.paragraphs[0]
        run = paragraph.add_run(category)
        run.font.size = Pt(10)
        
        # Center-align the text inside the cell
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

    set_table_borders(index_table)

    # Set alignment to center
    index_table.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Set cell height to 0.22 inches for the index table
    for row in index_table.rows:
        for cell in row.cells:
            tc = cell._tc
            tcPr = tc.get_or_add_tcPr()
            tcHeight = OxmlElement('w:trHeight')
            tcHeight.set(qn('w:val'), str(int(0.22 * 1440)))  # 0.22 inches in twips
            tcHeight.set(qn('w:hRule'), 'exact')
            tcPr.append(tcHeight)


    document.add_paragraph()  # Line break

    # Create the main data table
    create_granted_patents_data_table(document, df_grant, categories)


# Helper function to create the Granted Patents data table
def create_granted_patents_data_table(document, df_grant, categories):
    # Create main data table with 5 columns
    table = document.add_table(rows=1, cols=5)
    table.autofit = False  # Disable autofit to set custom widths

    # Set column widths
    column_widths = [Inches(0.4), Inches(1.08), Inches(2.46), Inches(1.38), Inches(1.48)]

    # Set column headers
    headers = ['Sl No', 'Patent No', 'Title', 'Assignee', 'Inventors']
    header_row = table.rows[0].cells

    for i, (header, width) in enumerate(zip(headers, column_widths)):
        cell = header_row[i]
        cell.text = header
        cell.width = width
        paragraph = cell.paragraphs[0]
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        paragraph.runs[0].bold = True
        paragraph.runs[0].font.size = Pt(10)

    # Process each category and add corresponding data
    for category in categories:
        # Add category heading row
        cat_row = table.add_row()
        cat_cell = cat_row.cells[0]
        cat_cell.merge(cat_row.cells[-1])
        cat_cell.text = f'> {category.upper()}'
        paragraph = cat_cell.paragraphs[0]
        paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
        paragraph.runs[0].bold = True
        paragraph.runs[0].font.size = Pt(10)

        # Filter data for the category
        cat_data = df_grant[df_grant['Category'].str.lower().str.contains(category.lower())]

        # Add rows for the data
        for _, row in cat_data.iterrows():
            data_row = table.add_row()
            data_row.cells[0].text = str(row.get('Serial No', ''))
            data_row.cells[1].text = str(row.get('Patent No', ''))
            data_row.cells[2].text = str(row.get('Title', ''))
            data_row.cells[3].text = str(row.get('Assignee', ''))
            data_row.cells[4].text = str(row.get('Inventors', ''))

            # Set column widths and font size for data rows
            for i, cell in enumerate(data_row.cells):
                cell.width = column_widths[i]
                for paragraph in cell.paragraphs:
                    paragraph.space_after = Pt(0)
                    paragraph.space_before = Pt(0)
                    for run in paragraph.runs:
                        run.font.size = Pt(10)

    set_table_borders(table)
# -----------------------------
# Part 5: Adding Detailed Publication Records
# -----------------------------

def add_detailed_publication_records_with_bookmarks(document, excel_path):
    df_fp = pd.read_excel(excel_path, sheet_name='FP')
    df_grant = pd.read_excel(excel_path, sheet_name='Grant')
    sheet1_df = pd.read_excel(excel_path, sheet_name='Sheet1')  # Load Sheet1 for images or other details
    
    headings = ['Serial No', 'Family number', 'Publication No', 'Kind Code', 'Title', 'Publication Date', 
                'Earliest Priority Date', 'Assignee', 'Inventors', 'Category', 'IPC', 'Patent Link', 'Abstract']

    # Pass sheet1_df to the function
    for _, row in df_fp.iterrows():
        add_detailed_table_with_bookmark(document, row, headings, sheet1_df)

    for _, row in df_grant.iterrows():
        add_detailed_table_with_bookmark(document, row, headings, sheet1_df)


# Helper function to add section headings
def add_section_heading(document, text):
    title = document.add_paragraph()
    title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    title_run = title.add_run(text)
    title_run.bold = True
    title_run.underline = True
    title_run.font.size = Pt(11)
    document.add_paragraph()  # Line break


# Helper function to process records and add detailed tables
def process_records(df, doc, sheet1_df, headings, section_title):
    doc.add_paragraph().add_run(section_title).bold = True
    
    for _, row in df.iterrows():
        row_data = row.to_dict()
        table = doc.add_table(rows=len(headings), cols=2)
        table.style = 'Table Grid'

        values = [
            format_date(str(row_data.get(field, ''))) if 'Date' in field else (
                str(row_data.get('Patent Link', '')) if field == 'PDF Document' else str(row_data.get(field, ''))
            )
            for field in headings[:-1]
        ]

        for i in range(len(headings) - 1):
            table.cell(i, 0).text = headings[i]
            table.cell(i, 1).text = values[i]

        image_title_cell = table.cell(len(headings) - 1, 0)
        image_title_cell.text = 'Image'
        image_title_cell.paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

        image_cell = table.cell(len(headings) - 1, 1)
        image_cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        image_cell.paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        image_cell.add_paragraph()  # Add space before image

        # Fetch image using Publication No instead of Family Number
        publication_no = row_data.get('Publication No', '')
        if publication_no and not sheet1_df[sheet1_df['Publication No'] == publication_no].empty:
            image_link = sheet1_df.loc[sheet1_df['Publication No'] == publication_no, 'Image'].values[0]
            try:
                response = requests.get(image_link)
                response.raise_for_status()
                image_stream = BytesIO(response.content)
                p = image_cell.add_paragraph()
                p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                run = p.add_run()
                run.add_picture(image_stream, width=Inches(3))
            except Exception as e:
                print(f"Error fetching image for Publication No {publication_no}: {e}")

        image_cell.add_paragraph()  # Add space after image
        doc.add_page_break()



# Helper function to add detailed records with bookmarks
def add_detailed_table_with_bookmark(document, row_data, headings, sheet1_df):
    publication_no = str(row_data.get('Publication No', row_data.get('Patent No', '')))
    bookmark_name = f'pub_{publication_no}'

    # Add a bookmark at the start of the detailed section
    paragraph = document.add_paragraph()
    run = paragraph.add_run()
    bookmark_start = OxmlElement('w:bookmarkStart')
    bookmark_start.set(qn('w:id'), '0')
    bookmark_start.set(qn('w:name'), bookmark_name)
    run._r.append(bookmark_start)

    # Add the detailed table
    add_detailed_table(document, row_data, headings)

    # Close the bookmark after the detailed section
    bookmark_end = OxmlElement('w:bookmarkEnd')
    bookmark_end.set(qn('w:id'), '0')
    run._r.append(bookmark_end)


# Helper to set the width of the table
def set_table_width(table, total_width_twips):
    tbl = table._tbl
    # Check if tblPr exists; if not, create it
    tblPr = tbl.find(qn('w:tblPr'))
    if tblPr is None:
        tblPr = OxmlElement('w:tblPr')
        tbl.insert(0, tblPr)

    # Set the table width
    tblW = OxmlElement('w:tblW')
    tblW.set(qn('w:w'), str(total_width_twips))
    tblW.set(qn('w:type'), 'dxa')
    tblPr.append(tblW)

    # Set the table layout to fixed
    tblLayout = OxmlElement('w:tblLayout')
    tblLayout.set(qn('w:type'), 'fixed')
    tblPr.append(tblLayout)



# Helper to add an image to the table cell
def add_image_to_cell(image_title_cell, image_cell, family_number, sheet1_df, right_width_twips):
    # Set the title for the image cell
    image_title_cell.text = 'Image'
    set_cell_width(image_title_cell, int(1.43 * 1440))
    set_cell_margins(image_title_cell)
    set_paragraph_font(image_title_cell.paragraphs[0])

    # Check if image is available for the Family number
    if family_number and not sheet1_df[sheet1_df['Family number'] == family_number].empty:
        image_link = sheet1_df.loc[sheet1_df['Family number'] == family_number, 'Image'].values[0]
        try:
            response = requests.get(image_link)
            response.raise_for_status()
            image_stream = BytesIO(response.content)

            # Add image to the cell
            p = image_cell.paragraphs[0]
            p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
            run = p.add_run()
            run.add_picture(image_stream, width=Inches(2))  # Adjust size if needed
        except Exception as e:
            print(f"Error fetching image for Family number {family_number}: {e}")

    set_cell_width(image_cell, right_width_twips)
    set_cell_margins(image_cell)
    image_cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
# -----------------------------
# Part 6: Final Integration and Saving the Document
# -----------------------------

def create_final_document(excel_path, template_path, output_path):
    # Initialize the document using the provided template
    document = Document(template_path)

    # 1. Add the First Two Pages (Patent Watch)
    add_first_two_pages(document, excel_path)
    document.add_page_break()  # Page break after the first two pages

    # 2. Add the First Publications Section
    add_first_publications_section(document, excel_path)
    document.add_page_break()  # Page break after First Publications

    # 3. Add the Granted Patents Section
    add_granted_patents_section(document, excel_path)
    document.add_page_break()  # Page break after Granted Patents

    # 4. Add Detailed Publication Records (First Publications & Granted Patents)
    add_detailed_publication_records_with_bookmarks(document, excel_path)

    # Save the final document
    document.save(output_path)
    print(f"Document successfully created: {output_path}")


# -----------------------------
# Usage Example
# -----------------------------

create_final_document(
    'C:/Users/Ayman/Documents/Abhijit_mail_attachments/Test_PW.xlsm',  # Excel file path
    'basic_page_template.docx',                                        # Template document
    'final_patent_watch.docx'                                          # Output document
)
