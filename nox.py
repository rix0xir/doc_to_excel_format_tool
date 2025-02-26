import pandas as pd
from docx import Document
from docx.shared import Inches, RGBColor, Pt
from docx.oxml.ns import qn, nsdecls
from docx.oxml import OxmlElement, parse_xml
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
import requests
from io import BytesIO
from PIL import Image
from datetime import datetime

# -------------------- Helper Functions -------------------- #

def resize_image(image_stream, fixed_height=4.5):
    img = Image.open(image_stream)
    original_width, original_height = img.size
    aspect_ratio = original_width / original_height
    new_height = int(fixed_height * 1440)
    new_width = int(new_height * aspect_ratio)
    resized_img = img.resize((new_width, new_height), Image.ANTIALIAS)
    return resized_img, new_width, new_height

def set_cell_margins(cell, top=20, start=20, bottom=20, end=20):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    for m in ('w:topMar', 'w:startMar', 'w:bottomMar', 'w:endMar'):
        for node in tcPr.findall(m, namespaces={'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'}):
            tcPr.remove(node)
    cell_margins = parse_xml(
        r'<w:tcMar %s>'
        r'<w:top w:w="%s" w:type="dxa"/>'
        r'<w:start w:w="%s" w:type="dxa"/>'
        r'<w:bottom w:w="%s" w:type="dxa"/>'
        r'<w:end w:w="%s" w:type="dxa"/>'
        r'</w:tcMar>' % (nsdecls('w'), top, start, bottom, end)
    )
    tcPr.append(cell_margins)

def set_cell_width(cell, width_twips):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    for child in tcPr.findall(qn('w:tcW')):
        tcPr.remove(child)
    tcW = OxmlElement('w:tcW')
    tcW.set(qn('w:w'), str(width_twips))
    tcW.set(qn('w:type'), 'dxa')
    tcPr.append(tcW)

def set_paragraph_font(paragraph, font_name='Calibri', font_size=Pt(10)):
    for run in paragraph.runs:
        run.font.name = font_name
        run.font.size = font_size

def add_blank_paragraph(cell):
    p = cell.paragraphs[0].insert_paragraph_before()
    p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    p = cell.add_paragraph()
    p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

def format_date(date_str):
    try:
        date_obj = datetime.strptime(date_str, '%Y-%m-%d %H:%M:%S')
        return date_obj.strftime('%d/%b/%Y')
    except ValueError:
        return date_str
    
# -------------------- Define Table Dimensions -------------------- #
left_width_twips = int(1.43 * 1440)
right_width_twips = int(5.06 * 1440)
total_width_twips = left_width_twips + right_width_twips

# -------------------- Load Excel Data -------------------- #
excel_path = 'C:/Users/Ayman/Documents/Abhijit_mail_attachments/Test_PW.xlsm'
first_pub_df = pd.read_excel(excel_path, sheet_name='First Publication')
granted_patents_df = pd.read_excel(excel_path, sheet_name='Granted')
sheet1_df = pd.read_excel(excel_path, sheet_name='Sheet1')

# -------------------- Open Document -------------------- #
doc_path = 'basic_page_template.docx'
doc = Document(doc_path)

def add_section_heading(doc, text):
    """Adds a centered, bold, and underlined section heading."""
    title = doc.add_paragraph()
    title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    title_run = title.add_run(text)
    title_run.bold = True
    title_run.underline = True
    title_run.font.size = Pt(11)

def add_index_link(doc):
    """Adds a right-aligned hyperlink labeled '<< INDEX'."""
    index_paragraph = doc.add_paragraph()
    index_paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
    index_run = index_paragraph.add_run('<< INDEX')
    index_run.font.color.rgb = RGBColor(0, 0, 255)
    index_run.font.underline = True
    hyperlink = OxmlElement('w:hyperlink')
    hyperlink.set(qn('r:id'), 'rId1')  # Make sure rId1 is properly referenced in the document relationships
    hyperlink.append(index_run._r)
    index_paragraph._element.append(hyperlink)

def add_first_index_template(doc):
    categories = [
        'Seafloor', 'Land', 'Marine', 'Microseismic & Multiphysics',
        'Processing', 'Reservoir', 'Geology', 'Data Management & Computing'
    ]

    # 1x8 Table with categories as hyperlinks
    table = doc.add_table(rows=1, cols=8)
    table.style = 'Table Grid'

    for i, category in enumerate(categories):
        cell = table.cell(0, i)
        paragraph = cell.paragraphs[0]
        run = paragraph.add_run(category)
        run.font.color.rgb = RGBColor(0, 0, 255)
        run.font.underline = True

        # Hyperlink placeholder (ensure hyperlinks are properly managed in your Word template)
        hyperlink = OxmlElement('w:hyperlink')
        hyperlink.set(qn('r:id'), 'rId1')  # Link ID must be correctly assigned in your document template
        hyperlink.append(run._r)
        paragraph._element.append(hyperlink)

    # Add 'First Publications' Heading
    title_paragraph = doc.add_paragraph()
    title_paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    title_run = title_paragraph.add_run('First Publications')
    title_run.bold = True
    title_run.font.size = Pt(12)
    doc.add_paragraph()  # Adds space below the heading

    # Column Headings for Publication Details
    headings = ['Sl No', 'Publication No', 'Title', 'Assignee', 'Inventors']
    pub_table = doc.add_table(rows=1, cols=len(headings))
    pub_table.style = 'Table Grid'

    # Set Column Widths
    set_cell_width(pub_table.cell(0, 0), int(0.75 * 1440))  # Sl No
    set_cell_width(pub_table.cell(0, 1), int(1.08 * 1440))  # Publication No
    set_cell_width(pub_table.cell(0, 2), int(2.46 * 1440))  # Title
    set_cell_width(pub_table.cell(0, 3), int(1.38 * 1440))  # Assignee
    set_cell_width(pub_table.cell(0, 4), int(1.48 * 1440))  # Inventors

    # Add Column Headings
    for i, heading in enumerate(headings):
        pub_table.cell(0, i).text = heading

    # Add Categories and their Data
    for category in categories:
        # Add Category Title (e.g., SEAFLOOR →)
        para = doc.add_paragraph()
        para.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
        run = para.add_run(f'{category.upper()} \u2192')  # Right arrow bullet
        run.bold = True

        # Filter and Add Data Rows from First Publications
        category_data_df = first_pub_df[first_pub_df['Category'] == category]
        for idx, row in category_data_df.iterrows():
            data_row = doc.add_table(rows=1, cols=len(headings))
            data_row.style = 'Table Grid'

            row_data = [str(idx + 1), str(row['Publication No']), row['Title'], row['Assignee'], row['Inventors']]
            for i, value in enumerate(row_data):
                data_row.cell(0, i).text = value

                # Add hyperlink to 'Publication No'
                if headings[i] == 'Publication No':
                    run = data_row.cell(0, i).paragraphs[0].runs[0]
                    run.font.color.rgb = RGBColor(0, 0, 255)
                    run.font.underline = True

                    hyperlink = OxmlElement('w:hyperlink')
                    hyperlink.set(qn('r:id'), 'rId1')  # Ensure this matches the correct reference in your doc
                    hyperlink.append(run._r)
                    data_row.cell(0, i).paragraphs[0]._element.append(hyperlink)

            doc.add_paragraph()  # Add space between rows

    doc.add_page_break()  # Move to the next section after First Publications Index

def add_second_index_template(doc):
    categories = [
        'Seafloor', 'Land', 'Marine', 'Microseismic & Multiphysics',
        'Processing', 'Reservoir', 'Geology', 'Data Management & Computing'
    ]

    # 1x8 Table with categories as hyperlinks
    table = doc.add_table(rows=1, cols=8)
    table.style = 'Table Grid'

    for i, category in enumerate(categories):
        cell = table.cell(0, i)
        paragraph = cell.paragraphs[0]
        run = paragraph.add_run(category)
        run.font.color.rgb = RGBColor(0, 0, 255)
        run.font.underline = True

        # Hyperlink placeholder (ensure hyperlinks are properly managed in your Word template)
        hyperlink = OxmlElement('w:hyperlink')
        hyperlink.set(qn('r:id'), 'rId2')  # Unique link ID for Granted Patents
        hyperlink.append(run._r)
        paragraph._element.append(hyperlink)

    # Add 'Granted Patents' Heading
    title_paragraph = doc.add_paragraph()
    title_paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    title_run = title_paragraph.add_run('Granted Patents')
    title_run.bold = True
    title_run.font.size = Pt(12)
    doc.add_paragraph()

    # Column Headings for Granted Patent Details
    headings = ['Sl No', 'Patent No', 'Title', 'Assignee', 'Inventors']
    patent_table = doc.add_table(rows=1, cols=len(headings))
    patent_table.style = 'Table Grid'

    # Set Column Widths
    set_cell_width(patent_table.cell(0, 0), int(0.75 * 1440))  # Sl No
    set_cell_width(patent_table.cell(0, 1), int(1.08 * 1440))  # Patent No
    set_cell_width(patent_table.cell(0, 2), int(2.46 * 1440))  # Title
    set_cell_width(patent_table.cell(0, 3), int(1.38 * 1440))  # Assignee
    set_cell_width(patent_table.cell(0, 4), int(1.48 * 1440))  # Inventors

    # Add Column Headings
    for i, heading in enumerate(headings):
        patent_table.cell(0, i).text = heading

    # Add Categories and their Data
    for category in categories:
        para = doc.add_paragraph()
        para.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
        run = para.add_run(f'{category.upper()} \u2192')
        run.bold = True

        # Filter and Add Data Rows from Granted Patents
        category_data_df = granted_patents_df[granted_patents_df['Category'] == category]
        for idx, row in category_data_df.iterrows():
            data_row = doc.add_table(rows=1, cols=len(headings))
            data_row.style = 'Table Grid'

            row_data = [str(idx + 1), str(row['Patent No']), row['Title'], row['Assignee'], row['Inventors']]
            for i, value in enumerate(row_data):
                data_row.cell(0, i).text = value

                # Add hyperlink to 'Patent No'
                if headings[i] == 'Patent No':
                    run = data_row.cell(0, i).paragraphs[0].runs[0]
                    run.font.color.rgb = RGBColor(0, 0, 255)
                    run.font.underline = True

                    hyperlink = OxmlElement('w:hyperlink')
                    hyperlink.set(qn('r:id'), 'rId2')
                    hyperlink.append(run._r)
                    data_row.cell(0, i).paragraphs[0]._element.append(hyperlink)

            doc.add_paragraph()

    doc.add_page_break()  # Move to the next section after Granted Patents Index

def process_records(df, headings, link_id):
    """
    Processes and adds detailed records for each publication or granted patent.
    Each record is placed on a separate page.
    """
    namespaces = {'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'}  # Namespace mapping

    for _, row in df.iterrows():
        row_data = row.to_dict()

        if doc.paragraphs[-1].text != '<< INDEX':
            add_index_link(doc)

        # Create a table with two columns (field name and value)
        table = doc.add_table(rows=len(headings), cols=2)
        table.style = 'Table Grid'

        tbl = table._tbl
        tblPr = tbl.find('./w:tblPr', namespaces=namespaces)
        if tblPr is None:
            tblPr = OxmlElement('w:tblPr')
            tbl.insert(0, tblPr)
        tblW = OxmlElement('w:tblW')
        tblW.set(qn('w:w'), str(total_width_twips))
        tblW.set(qn('w:type'), 'dxa')
        tblPr.append(tblW)
        tblLayout = OxmlElement('w:tblLayout')
        tblLayout.set(qn('w:type'), 'fixed')
        tblPr.append(tblLayout)

        # Populate the table with data
        for i, field in enumerate(headings[:-1]):  # Exclude the 'Image' field
            cell_left = table.cell(i, 0)
            cell_right = table.cell(i, 1)

            # Left column: Field name
            cell_left.text = field
            set_cell_width(cell_left, left_width_twips)
            set_cell_margins(cell_left)
            set_paragraph_font(cell_left.paragraphs[0])

            # Right column: Field value
            value = format_date(str(row_data.get(field, ''))) if 'Date' in field else str(row_data.get(field, ''))
            cell_right.text = value
            set_cell_width(cell_right, right_width_twips)
            set_cell_margins(cell_right)
            set_paragraph_font(cell_right.paragraphs[0])

            # Add hyperlink to Publication/Patent No
            if field in ['Publication No', 'Patent No']:
                run = cell_right.paragraphs[0].runs[0]
                run.font.color.rgb = RGBColor(0, 0, 255)
                run.font.underline = True

                hyperlink = OxmlElement('w:hyperlink')
                hyperlink.set(qn('r:id'), link_id)  # Use the correct link ID
                hyperlink.append(run._r)
                cell_right.paragraphs[0]._element.append(hyperlink)

            # Image Handling
            image_title_cell = table.cell(len(headings) - 1, 0)
            image_title_cell.text = 'Image'
            set_cell_width(image_title_cell, left_width_twips)
            set_cell_margins(image_title_cell)
            set_paragraph_font(image_title_cell.paragraphs[0])

            image_cell = table.cell(len(headings) - 1, 1)
            image_cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            set_cell_width(image_cell, right_width_twips)
            set_cell_margins(image_cell)

            # Fetch image from Sheet1 using Family number
            family_number = row_data.get('Family number', '')
            if family_number and not sheet1_df[sheet1_df['Family number'] == family_number].empty:
                image_link = sheet1_df.loc[sheet1_df['Family number'] == family_number, 'Image'].values[0]
                try:
                    response = requests.get(image_link)
                    response.raise_for_status()
                    image_stream = BytesIO(response.content)
                    p = image_cell.paragraphs[0]
                    p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                    run = p.add_run()
                    run.add_picture(image_stream, width=Inches(2))
                except Exception as e:
                    print(f"Error fetching image for Family number {family_number}: {e}")

            add_blank_paragraph(image_cell)
            doc.add_page_break()


# Headings for First Publications
first_pub_headings = [
    'Serial No', 'Family number', 'Publication No', 'Kind Code', 'Title', 'Publication Date',
    'Earliest Priority Date', 'Assignee', 'Inventors', 'Category', 'IPC',
    'Patent Link', 'Abstract', 'Image'
]

# Headings for Granted Patents
granted_patent_headings = [
    'Serial No', 'Family number', 'Patent No', 'Kind Code', 'Title', 'Publication Date',
    'Earliest Priority Date', 'Assignee', 'Inventors', 'Category', 'IPC',
    'Patent Link', 'Abstract', 'Image'
]

# Add First Index Page
add_first_index_template(doc)

# Add Second Index Page
add_second_index_template(doc)

# Add First Publications Section
add_section_heading(doc, 'FIRST PUBLICATIONS')
add_index_link(doc)
process_records(first_pub_df, first_pub_headings, link_id='rId1')

# Add Granted Patents Section
add_section_heading(doc, 'GRANTED PATENTS')
add_index_link(doc)
process_records(granted_patents_df, granted_patent_headings, link_id='rId2')

# Save the Document
doc.save('updated_publications.docx')
print("Document successfully created: 'updated_publications.docx'")
