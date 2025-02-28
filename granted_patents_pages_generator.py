import pandas as pd
from docx import Document
from docx.shared import Pt, Inches
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.oxml.shared import OxmlElement
import sys
import logging
import requests
from io import BytesIO
from PIL import Image
import os
import datetime

# Set up logging
logging.basicConfig(level=logging.INFO, format='%(asctime)s - %(levelname)s - %(message)s')
logger = logging.getLogger(__name__)

class PageTracker:
    """Enhanced page tracking with element measurement"""
    def __init__(self, document):
        self.document = document
        self.current_page_height = 0.0
        self.page_height = 11.0  # Letter size
        self.margin_top = document.sections[0].top_margin.inches
        self.margin_bottom = document.sections[0].bottom_margin.inches
        self.usable_height = self.page_height - self.margin_top - self.margin_bottom
        self.header_height = 0.4  # Measured height of <<INDEX header
        
    def check_space(self, required_height):
        """Check if required height fits current page"""
        remaining = self.usable_height - self.current_page_height
        return remaining >= required_height
        
    def add_page_break(self):
        """Create new page with proper header accounting"""
        self.document.add_page_break()
        # Add header with exact height measurement
        header_para = self.document.add_paragraph("<<INDEX")
        header_para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        header_para.paragraph_format.space_after = Pt(0)
        header_para.paragraph_format.space_before = Pt(0)
        header_para.runs[0].font.size = Pt(10)
        self.current_page_height = self.header_height

def estimate_table_height(table_data, has_image):
    """Calculate approximate table height in inches"""
    base_height = len(table_data) * 0.28  # Regular rows
    abstract_text = next((v for (k, v) in table_data if k == "Abstract"), "")
    abstract_lines = max(1, len(str(abstract_text)) // 100)  # 100 chars/line
    abstract_height = abstract_lines * 0.2
    image_height = 3.0 if has_image else 0.0  # Max image height
    return base_height + abstract_height + image_height + 0.2  # Padding

def add_hyperlink(paragraph, text, url):
    """Add a hyperlink to a paragraph."""
    part = paragraph.part
    r_id = part.relate_to(url, 'http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink', is_external=True)
    
    hyperlink = OxmlElement('w:hyperlink')
    hyperlink.set(qn('r:id'), r_id)
    
    run = OxmlElement('w:r')
    rPr = OxmlElement('w:rPr')
    
    underline = OxmlElement('w:u')
    underline.set(qn('w:val'), 'single')
    rPr.append(underline)
    
    color = OxmlElement('w:color')
    color.set(qn('w:val'), '0000FF')
    rPr.append(color)
    
    run.append(rPr)
    t = OxmlElement('w:t')
    t.text = text
    run.append(t)
    hyperlink.append(run)
    paragraph._element.append(hyperlink)

def download_image(image_url, folder_path, family_number):
    """Download and save image."""
    try:
        response = requests.get(image_url, stream=True)
        response.raise_for_status()
        
        filename = f"patent_{family_number}.png"
        img_path = os.path.join(folder_path, filename)
        
        img = Image.open(BytesIO(response.content))
        img.save(img_path, format='PNG')
        return img_path
    except Exception as e:
        logger.error(f"Failed to download image: {e}")
        return None

def insert_image(cell, img_path):
    """Insert image into cell maintaining aspect ratio."""
    try:
        img = Image.open(img_path)
        aspect_ratio = img.height / img.width
        
        # Target width is cell width (5.61 inches)
        target_width = 5.61
        target_height = target_width * aspect_ratio
        
        # If height exceeds 3 inches, scale down
        if target_height > 3:
            target_height = 3
            target_width = target_height / aspect_ratio
        
        paragraph = cell.add_paragraph()
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        paragraph.paragraph_format.space_before = Pt(0)
        paragraph.paragraph_format.space_after = Pt(0)
        run = paragraph.add_run()
        run.add_picture(img_path, width=Inches(target_width), height=Inches(target_height))
        return True
    except Exception as e:
        logger.error(f"Error inserting image: {e}")
        return False

def create_patent_table(document, record, df_images, folder_path, page_tracker):
    """Create table with proactive page management"""
    # Define table data
    table_data = [
        ("Serial No", record.get('Serial No', "")),
        ("Family number", record.get('Family number', "")),
        ("Patent No", record.get('Patent No', "")),
        ("Kind Code", record.get('Kind Code', "")),
        ("Title", record.get('Title', "")),
        ("Publication Date", record.get('Publication Date', "")),
        ("Earliest Priority", record.get('Earliest Priority', "")),
        ("Assignee", record.get('Assignee', "")),
        ("Inventors", record.get('Inventors', "")),
        ("Category", record.get('Category', "")),
        ("IPC", record.get('IPC', "")),
        ("Patent Link", record.get('Patent Link', "")),
        ("Abstract", record.get('Abstract', ""))
    ]
    
    # Check if record has image
    has_image = False
    if pd.notna(record.get("Family number")):
        image_row = df_images[df_images['Family number'] == record['Family number']]
        has_image = not image_row.empty and pd.notna(image_row.iloc[0].get('Image'))
    
    # Calculate required space
    required_height = estimate_table_height(table_data, has_image)
    
    # Check space and add page break if needed
    if not page_tracker.check_space(required_height):
        page_tracker.add_page_break()
    
    # Create table
    table = document.add_table(rows=0, cols=2)
    table.style = 'Table Grid'
    table.columns[0].width = Inches(1.38)
    table.columns[1].width = Inches(5.61)
    
    # Add content
    for field_name, field_value in table_data:
        row_cells = table.add_row().cells
        row_cells[0].text = field_name
        row_cells[0].vertical_alignment = WD_ALIGN_VERTICAL.TOP
        row_cells[1].text = str(field_value) if pd.notna(field_value) else ""
        row_cells[1].vertical_alignment = WD_ALIGN_VERTICAL.TOP
        
        if field_name == "Patent Link" and pd.notna(field_value):
            pdf_para = row_cells[1].paragraphs[0]
            pdf_para.clear()
            add_hyperlink(pdf_para, "Link", field_value)
    
    # Handle image if exists
    if has_image:
        img_path = download_image(image_row.iloc[0]['Image'], folder_path, record['Family number'])
        if img_path:
            image_cells = table.add_row().cells
            image_cells[0].text = "Image"
            image_cells[0].vertical_alignment = WD_ALIGN_VERTICAL.TOP
            insert_image(image_cells[1], img_path)
    
    # Update page tracker
    page_tracker.current_page_height += required_height
    return table

def create_granted_patents_document(document, df_granted, df_images, folder_path):
    """Main document creation flow with proper page tracking"""
    page_tracker = PageTracker(document)
    
    # Add heading
    heading = document.add_paragraph("GRANTED PATENTS")
    heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
    heading.runs[0].bold = True
    heading.runs[0].underline = True
    heading.runs[0].font.size = Pt(11)
    heading.paragraph_format.space_before = Pt(0)
    heading.paragraph_format.space_after = Pt(0)
    page_tracker.current_page_height += 0.6  # Measured header height
    
    # Add initial index
    index = document.add_paragraph("<<INDEX")
    index.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    index.runs[0].font.size = Pt(10)
    index.paragraph_format.space_before = Pt(0)
    index.paragraph_format.space_after = Pt(0)
    page_tracker.current_page_height += page_tracker.header_height
    
    # Process each record
    for idx, row in df_granted.iterrows():
        if idx > 0 and not page_tracker.check_space(0.3):  # Check space for new record
            page_tracker.add_page_break()
        
        create_patent_table(document, row.to_dict(), df_images, folder_path, page_tracker)

def main():
    try:
        excel_path = sys.argv[1] if len(sys.argv) > 1 else r'C:\Users\Ayman\Documents\Abhijit_mail_attachments\Test_PW.xlsm'
        df_granted = pd.read_excel(excel_path, sheet_name="Granted")
        df_images = pd.read_excel(excel_path, sheet_name="Sheet1")
        
        document = Document("basic_page_template.docx")
        
        # Create image folder
        image_folder = f"patent_images_{datetime.datetime.now().strftime('%Y%m%d_%H%M%S')}"
        os.makedirs(image_folder, exist_ok=True)
        
        create_granted_patents_document(document, df_granted, df_images, image_folder)
        
        document.save("part_5.docx")
        print("Document created successfully")
        
    except Exception as e:
        print(f"Error: {e}")
        raise

if __name__ == "__main__":
    main()