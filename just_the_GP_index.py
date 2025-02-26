import pandas as pd
from docx import Document
from docx.shared import Inches, Pt
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ROW_HEIGHT_RULE
import sys
from docx.shared import RGBColor

# Function to set table borders
def set_table_borders(table):
    tbl = table._tbl
    tblPr = tbl.tblPr
    borders = OxmlElement('w:tblBorders')
    for border_name in ['top', 'left', 'bottom', 'right', 'insideH', 'insideV']:
        border = OxmlElement(f'w:{border_name}')
        border.set(qn('w:val'), 'single')
        border.set(qn('w:sz'), '4')  # Size of the border
        border.set(qn('w:space'), '0')
        border.set(qn('w:color'), '000000')  # Black color for the border
        borders.append(border)
    tblPr.append(borders)


def add_hyperlink(paragraph, text, bookmark_name):
    """Add an internal hyperlink (bookmark link) in a Word document."""
    run = paragraph.add_run(text)
    run.font.color.rgb = RGBColor(0, 0, 255)  # Blue color for hyperlink
    run.underline = True

    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("w:anchor"), bookmark_name)  # Reference the bookmark name
    hyperlink.append(run._r)
    paragraph._element.append(hyperlink)


def create_granted_patents_doc(excel_path, output_path, template_path):
    # Read Excel data from the 'Grant' worksheet
    df = pd.read_excel(excel_path, sheet_name='Grant')
    
    # Clean column names and ensure there are no extra spaces
    df.columns = df.columns.str.strip()
    
    # Clean the 'Category' column values
    df['Category'] = df['Category'].astype(str).str.strip()

    # Load the template document
    document = Document(template_path)

    # Define categories, including missing ones
    categories = [
        'Seafloor', 'Land', 'Marine', 'Microseismic & Multiphysics',
        'Processing', 'Reservoir', 'Geology', 'Data Management & Computing',
        'Downhole'
    ]

    # Create index table with placeholders
    index_table = document.add_table(rows=1, cols=len(categories))
    index_row = index_table.rows[0].cells
    for i, category in enumerate(categories):
        paragraph = index_row[i].paragraphs[0]
        run = paragraph.add_run(category)
        run.font.size = Pt(10)
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Set table borders for index table
    set_table_borders(index_table)

    document.add_paragraph()  # Line break

    # Add main heading
    heading_paragraph = document.add_paragraph('GRANTED PATENTS')
    heading_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    heading_run = heading_paragraph.runs[0]
    heading_run.bold = True
    heading_run.font.size = Pt(12)

    # Create main data table with specified column widths
    table = document.add_table(rows=1, cols=5)
    table.autofit = False  # Disable autofit to manually set column widths

    # Set column widths
    column_widths = [Inches(0.4), Inches(1.08), Inches(2.46), Inches(1.38), Inches(1.48)]

    # Set column headers
    headers = ['Sl No', 'Patent No', 'Title', 'Assignee', 'Inventors']
    header_row = table.rows[0].cells
    for i, (header, width) in enumerate(zip(headers, column_widths)):
        cell = header_row[i]
        cell.text = header
        cell.width = width
        paragraph = cell.paragraphs[0]
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        paragraph.runs[0].bold = True
        paragraph.runs[0].font.size = Pt(10)
        paragraph.space_after = Pt(0)
        paragraph.space_before = Pt(0)

    # Process each category
    for category in categories:
        # Add category heading row
        cat_row = table.add_row()
        cat_row.height = Inches(0.24)  # Set the row height to 0.27 inches
        cat_row.height_rule = WD_ROW_HEIGHT_RULE.EXACTLY  # Ensures the row height is fixed

        cat_cell = cat_row.cells[0]
        cat_cell.merge(cat_row.cells[-1])
        cat_cell.text = f'> {category.upper()}'
        paragraph = cat_cell.paragraphs[0]
        paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
        paragraph.runs[0].bold = True
        paragraph.runs[0].font.size = Pt(10)

        # Filter the DataFrame for rows matching the current category (case-insensitive)
        cat_data = df[df['Category'].str.lower().str.contains(category.lower())]

        # Add rows for the data
        for _, row in cat_data.iterrows():
            data_row = table.add_row()
            data_row.cells[0].text = str(row.get('Serial No', ''))

            pat_no = str(row.get('Patent No', ''))
            cell = data_row.cells[1]  # Assuming Patent No is in the second column
            
            if pat_no:
                p = cell.paragraphs[0]
                add_hyperlink(p, pat_no, pat_no.strip()) # The text and bookmark name are the same
            else:
                cell.text = pat_no  # Just display the text if no number exists

            data_row.cells[2].text = str(row.get('Title', ''))
            data_row.cells[3].text = str(row.get('Assignee', ''))
            data_row.cells[4].text = str(row.get('Inventors', ''))

            # Set column widths and font size for data rows
            for i, cell in enumerate(data_row.cells):
                cell.width = column_widths[i]
                for paragraph in cell.paragraphs:
                    paragraph.space_after = Pt(0)
                    paragraph.space_before = Pt(0)
                    for run in paragraph.runs:
                        run.font.size = Pt(10)

    # Set table borders for main data table
    set_table_borders(table)

    # Save document
    document.save(output_path)

# # Usage example:
# create_granted_patents_doc(
#     'C:/Users/Ayman/Documents/Abhijit_mail_attachments/Test_PW.xlsm',
#     'part_3.docx',
#     'basic_page_template.docx'
# )



# if __name__ == "__main__":
#     excel_path = sys.argv[1]
#     output_file = sys.argv[2]
#     template_file = sys.argv[3]

#     create_granted_patents_doc(excel_path, output_file, template_file)

# Run the function directly to generate output
excel_path = r"C:\Users\Ayman\Documents\Abhijit_mail_attachments\Test_PW.xlsm"
output_file = "GP_output.docx"
template_file = "basic_page_template.docx"

create_granted_patents_doc(excel_path, output_file, template_file)

print(f"Granted Patents index has been generated: {output_file}")
