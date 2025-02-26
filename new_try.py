debug this code(don't change logic, just if anything is out of place, repeated, without imports, etc:

import pandas as pd
from docx import Document
from docx.shared import Inches, RGBColor, Pt
from docx.oxml.ns import qn, nsdecls
from docx.oxml import OxmlElement, parse_xml
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
import requests
from io import BytesIO
from PIL import Image
from datetime import datetime

# --------------------
# Helper functions
# --------------------

def resize_image(image_stream, fixed_height=4.5):
    img = Image.open(image_stream)
    original_width, original_height = img.size
    aspect_ratio = original_width / original_height
    new_height = int(fixed_height * 1440)
    new_width = int(new_height * aspect_ratio)
    resized_img = img.resize((new_width, new_height), Image.ANTIALIAS)
    return resized_img, new_width, new_height

def set_cell_margins(cell, top=20, start=20, bottom=20, end=20):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    for m in ('w:topMar', 'w:startMar', 'w:bottomMar', 'w:endMar'):
        for node in tcPr.findall(m, namespaces={'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'}):
            tcPr.remove(node)
    cell_margins = parse_xml(
        r'<w:tcMar %s>'
        r'<w:top w:w="%s" w:type="dxa"/>' 
        r'<w:start w:w="%s" w:type="dxa"/>' 
        r'<w:bottom w:w="%s" w:type="dxa"/>' 
        r'<w:end w:w="%s" w:type="dxa"/>'
        r'</w:tcMar>' % (nsdecls('w'), top, start, bottom, end)
    )
    tcPr.append(cell_margins)

def set_cell_width(cell, width_twips):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    for child in tcPr.findall(qn('w:tcW')):
        tcPr.remove(child)
    tcW = OxmlElement('w:tcW')
    tcW.set(qn('w:w'), str(width_twips))
    tcW.set(qn('w:type'), 'dxa')
    tcPr.append(tcW)

def set_paragraph_font(paragraph, font_name='Calibri', font_size=Pt(10)):
    for run in paragraph.runs:
        run.font.name = font_name
        run.font.size = font_size

def add_blank_paragraph(cell):
    p = cell.paragraphs[0].insert_paragraph_before()
    p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    p = cell.add_paragraph()
    p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER


def format_date(date_str):
    # Assuming the date_str is in 'YYYY-MM-DD HH:MM:SS' format
    try:
        date_obj = datetime.strptime(date_str, '%Y-%m-%d %H:%M:%S')
        return date_obj.strftime('%d/%b/%Y')
    except ValueError:
        return date_str

# --------------------
# Define table dimensions
# --------------------
left_width_twips = int(1.43 * 1440)
right_width_twips = int(5.06 * 1440)
total_width_twips = left_width_twips + right_width_twips

# --------------------
# Load Excel Data
# --------------------
excel_path = 'C:/Users/Ayman/Documents/Abhijit_mail_attachments/Test_PW.xlsm'
first_pub_df = pd.read_excel(excel_path, sheet_name='First Publication')
granted_patents_df = pd.read_excel(excel_path, sheet_name='Granted')
sheet1_df = pd.read_excel(excel_path, sheet_name='Sheet1')

# --------------------
# Open Document and add Title/Index
# --------------------
doc_path = 'basic_page_template.docx'
doc = Document(doc_path)

def add_section_heading(doc, text):
    title = doc.add_paragraph()
    title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    title_run = title.add_run(text)
    title_run.bold = True
    title_run.underline = True
    title_run.font.size = Pt(11)
    
def add_index_link(doc):
    index_paragraph = doc.add_paragraph()
    index_paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
    index_run = index_paragraph.add_run('<< INDEX')
    index_run.font.color.rgb = RGBColor(0, 0, 255)
    index_run.font.underline = True
    hyperlink = OxmlElement('w:hyperlink')
    hyperlink.set(qn('r:id'), 'rId1')
    hyperlink.append(index_run._r)
    index_paragraph._element.append(hyperlink)

def add_first_index_template(doc):
    # 1x8 Table with categories as hyperlinks
    categories = [
        'Seafloor', 'Land', 'Marine', 'Microseismic & Multiphysics', 
        'Processing', 'Reservoir', 'Geology', 'Data Management & Computing'
    ]
    
    # Add table
    table = doc.add_table(rows=1, cols=8)
    table.style = 'Table Grid'
    for i, category in enumerate(categories):
        cell = table.cell(0, i)
        paragraph = cell.paragraphs[0]
        run = paragraph.add_run(category)
        run.font.color.rgb = RGBColor(0, 0, 255)
        run.font.underline = True
        # Add hyperlink functionality
        hyperlink = OxmlElement('w:hyperlink')
        hyperlink.set(qn('r:id'), 'rId1')
        hyperlink.append(run._r)
        paragraph._element.append(hyperlink)

    # Add "First Publications" heading
    title_paragraph = doc.add_paragraph()
    title_paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    title_run = title_paragraph.add_run('First Publications')
    title_run.bold = True
    title_run.font.size = Pt(12)
    
    doc.add_paragraph()  # Empty line between title and table

    # --------------------
    # Create Table for Publication Details
    # --------------------
    headings = ['Sl No', 'Publication No', 'Title', 'Assignee', 'Inventors']
    table = doc.add_table(rows=1, cols=len(headings))
    table.style = 'Table Grid'

    # Set column widths
    set_cell_width(table.cell(0, 0), left_width_twips)  # Sl No
    set_cell_width(table.cell(0, 1), int(1.08 * 1440))  # Publication No
    set_cell_width(table.cell(0, 2), int(2.46 * 1440))  # Title
    set_cell_width(table.cell(0, 3), int(1.38 * 1440))  # Assignee
    set_cell_width(table.cell(0, 4), int(1.48 * 1440))  # Inventors

    # Add headings to the first row
    for i, heading in enumerate(headings):
        table.cell(0, i).text = heading

    # --------------------
    # Extract and Add Category Data (Seafloor, Land, etc.)
    # --------------------
    categories_data = {
        'Seafloor': 'Category 1 data from FP',  # Placeholder, update with actual data extraction logic
        'Land': 'Category 2 data from FP',
        'Marine': 'Category 3 data from FP',
        'Microseismic & Multiphysics': 'Category 4 data from FP',
        'Processing': 'Category 5 data from FP',
        'Reservoir': 'Category 6 data from FP',
        'Geology': 'Category 7 data from FP',
        'Data Management & Computing': 'Category 8 data from FP'
    }

    # Iterate through categories and add bullet-pointed rows
    for category, data in categories_data.items():
        # Add category title with bullet
        para = doc.add_paragraph()
        para.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
        run = para.add_run(f'{category} \u2192')  # Arrow bullet (right arrow)
        run.bold = True
        
        # Add corresponding data rows for each category
        category_data_df = first_pub_df[first_pub_df['Category'] == category]
        for idx, row in category_data_df.iterrows():
            table = doc.add_table(rows=1, cols=len(headings))
            row_data = [str(row[col]) for col in headings[1:]]  # Skip 'Sl No' column
            for i, data_value in enumerate(row_data):
                table.cell(0, i).text = data_value
            doc.add_paragraph()  # Empty line between records
    doc.add_page_break()


 Add First Publications Section
add_section_heading(doc, 'FIRST PUBLICATIONS')
add_index_link(doc)

def process_records(df, headings):
    nsmap = {'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'}  # Add this line

    for _, row in df.iterrows():
        row_data = row.to_dict()

        if doc.paragraphs[-1].text != '<< INDEX':
            add_index_link(doc)

        table = doc.add_table(rows=len(headings), cols=2)
        table.style = 'Table Grid'

        tbl = table._tbl
        tblPr = tbl.find('./w:tblPr', namespaces=nsmap)
        if tblPr is None:
            tblPr = OxmlElement('w:tblPr')
            tbl.insert(0, tblPr)
        tblW = OxmlElement('w:tblW')
        tblW.set(qn('w:w'), str(total_width_twips))
        tblW.set(qn('w:type'), 'dxa')
        tblPr.append(tblW)
        tblLayout = OxmlElement('w:tblLayout')
        tblLayout.set(qn('w:type'), 'fixed')
        tblPr.append(tblLayout)

        values = [
            format_date(str(row_data.get(field, ''))) if 'Date' in field else str(row_data.get(field, ''))
            for field in headings[:-1]
        ]


        for i in range(len(headings) - 1):
            cell_left = table.cell(i, 0)
            cell_right = table.cell(i, 1)
            cell_left.text = headings[i]
            cell_right.text = values[i]

            set_cell_width(cell_left, left_width_twips)
            set_cell_width(cell_right, right_width_twips)
            set_cell_margins(cell_left)
            set_cell_margins(cell_right)
            for paragraph in cell_left.paragraphs + cell_right.paragraphs:
                set_paragraph_font(paragraph)

        image_title_cell = table.cell(len(headings) - 1, 0)
        image_title_cell.text = 'Image'
        set_cell_width(image_title_cell, left_width_twips)
        set_cell_margins(image_title_cell)
        for paragraph in image_title_cell.paragraphs:
            set_paragraph_font(paragraph)

        image_cell = table.cell(len(headings) - 1, 1)
        image_cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        set_cell_width(image_cell, right_width_twips)
        set_cell_margins(image_cell)

        family_number = row_data.get('Family number', '')
        if family_number and not sheet1_df[sheet1_df['Family number'] == family_number].empty:
            image_link = sheet1_df.loc[sheet1_df['Family number'] == family_number, 'Image'].values[0]
            try:
                response = requests.get(image_link)
                response.raise_for_status()
                image_stream = BytesIO(response.content)
                p = image_cell.paragraphs[0]
                p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                run = p.add_run()
                run.add_picture(image_stream, width=Inches(2))
            except Exception as e:
                print(f"Error fetching image for Family number {family_number}: {e}")

        add_blank_paragraph(image_cell)

        doc.add_page_break()

# Process First Publications
first_pub_headings = [
    'Serial No', 'Family number','Publication No', 'Kind Code', 'Title', 'Publication Date',
    'Earliest Priority Date', 'Assignee', 'Inventors', 'Category', 'IPC',
    'Patent Link', 'Abstract', 'Image'
]
process_records(first_pub_df, first_pub_headings)


def add_second_index_template(doc):
    # 1x8 Table with categories as hyperlinks for granted patents
    categories = [
        'Seafloor', 'Land', 'Marine', 'Microseismic & Multiphysics', 
        'Processing', 'Reservoir', 'Geology', 'Data Management & Computing'
    ]
    
    # Add table
    table = doc.add_table(rows=1, cols=8)
    table.style = 'Table Grid'
    for i, category in enumerate(categories):
        cell = table.cell(0, i)
        paragraph = cell.paragraphs[0]
        run = paragraph.add_run(category)
        run.font.color.rgb = RGBColor(0, 0, 255)
        run.font.underline = True
        # Add hyperlink functionality
        hyperlink = OxmlElement('w:hyperlink')
        hyperlink.set(qn('r:id'), 'rId2')
        hyperlink.append(run._r)
        paragraph._element.append(hyperlink)

    # Add "Granted Patents" heading
    title_paragraph = doc.add_paragraph()
    title_paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    title_run = title_paragraph.add_run('Granted Patents')
    title_run.bold = True
    title_run.font.size = Pt(12)
    
    doc.add_paragraph()  # Empty line between title and table

    # --------------------
    # Extract and Add Category Data (Seafloor, Land, etc. for Granted Patents)
    # --------------------
    categories_data = {
        'Seafloor': 'Category 1 data from Grant',  # Placeholder, update with actual data extraction logic
        'Land': 'Category 2 data from Grant',
        'Marine': 'Category 3 data from Grant',
        'Microseismic & Multiphysics': 'Category 4 data from Grant',
        'Processing': 'Category 5 data from Grant',
        'Reservoir': 'Category 6 data from Grant',
        'Geology': 'Category 7 data from Grant',
        'Data Management & Computing': 'Category 8 data from Grant'
    }

    # Iterate through categories and add bullet-pointed rows
    for category, data in categories_data.items():
        # Add category title with bullet
        para = doc.add_paragraph()
        para.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
        run = para.add_run(f'{category} \u2192')  # Arrow bullet (right arrow)
        run.bold = True
        
        # Add corresponding data rows for each category
        category_data_df = granted_patents_df[granted_patents_df['Category'] == category]
        for idx, row in category_data_df.iterrows():
            table = doc.add_table(rows=1, cols=len(headings))
            row_data = [str(row[col]) for col in headings[1:]]  # Skip 'Sl No' column
            for i, data_value in enumerate(row_data):
                table.cell(0, i).text = data_value
            doc.add_paragraph()  # Empty line between records
    doc.add_page_break()


# Add Granted Patents Section
add_section_heading(doc, 'GRANTED PATENTS')
add_index_link(doc)

# Process Granted Patents
granted_patent_headings = [
    'Serial No', 'Family number', 'Patent No', 'Kind Code', 'Title', 'Publication Date',
    'Earliest Priority Date', 'Assignee', 'Inventors', 'Category', 'IPC',
    'Patent Link', 'Abstract', 'Image'
]
process_records(granted_patents_df, granted_patent_headings)

add_first_index_template(doc)

# Add the First Publications Section
add_section_heading(doc, 'FIRST PUBLICATIONS')
add_index_link(doc)

# Process First Publications
process_records(first_pub_df, first_pub_headings)

# Add the Second Index Template
add_second_index_template(doc)

# Add the Granted Patents Section
add_section_heading(doc, 'GRANTED PATENTS')
add_index_link(doc)

# Process Granted Patents
process_records(granted_patents_df, granted_patent_headings)

# Save the document
doc.save('updated_publications.docx')
print("Document successfully created: 'updated_publications.docx'")