import pandas as pd
from docx import Document
from docx.shared import Pt, Inches
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_UNDERLINE
from docx.enum.table import WD_ALIGN_VERTICAL
import sys
import math

# ------------------------------
# Helper Functions
# ------------------------------

def set_table_borders(table):
    """Apply black borders to all sides of a table."""
    tbl = table._tbl
    tblPr = tbl.tblPr
    borders = OxmlElement('w:tblBorders')

    for border_name in ['top', 'left', 'bottom', 'right', 'insideH', 'insideV']:
        border = OxmlElement(f'w:{border_name}')
        border.set(qn('w:val'), 'single')
        border.set(qn('w:sz'), '4')  # Border thickness
        border.set(qn('w:space'), '0')  # No spacing
        border.set(qn('w:color'), '000000')  # Black color
        borders.append(border)

    tblPr.append(borders)


def set_row_height_exact(row, height_in_inches):
    """Set a fixed height for a table row."""
    tr = row._tr
    trPr = tr.get_or_add_trPr()
    trHeight = OxmlElement('w:trHeight')
    trHeight.set(qn('w:val'), str(int(height_in_inches * 1440)))  # Convert inches to twips
    trHeight.set(qn('w:hRule'), 'exact')  # Enforce exact height
    trPr.append(trHeight)


def set_cell_size_and_alignment(table, cell_height=0.28, cell_width=Inches(1.38)):
    """Set cell width, height, and alignment for all table cells."""
    for row in table.rows:
        set_row_height_exact(row, cell_height)  # Ensure consistent row height
        for cell in row.cells:
            cell.width = cell_width
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            for paragraph in cell.paragraphs:
                paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    table.alignment = WD_ALIGN_PARAGRAPH.RIGHT


# ------------------------------
# Document Content Functions
# ------------------------------

def create_company_table(document, first_row_data, second_row_data):
    """Create a company table with a fixed row height of 0.37 inches."""
    table = document.add_table(rows=2, cols=6)

    # Fill first row
    for i, company in enumerate(first_row_data):
        cell = table.rows[0].cells[i]
        cell.text = company
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        cell.paragraphs[0].runs[0].font.size = Pt(10)

    # Merge last two cells in first row if data is <6 items
    if len(first_row_data) < 6:
        table.rows[0].cells[-2].merge(table.rows[0].cells[-1])

    # Fill second row
    for i, company in enumerate(second_row_data):
        cell = table.rows[1].cells[i]
        cell.text = company
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        cell.paragraphs[0].runs[0].font.size = Pt(10)

    # Apply styles
    set_table_borders(table)
    set_cell_size_and_alignment(table, cell_height=0.37)  # Fixed height for company table
    table.alignment = WD_ALIGN_PARAGRAPH.CENTER
    table.autofit = True

    return table


def add_patent_section_links(document):
    """Add First Publications and Granted Patents hyperlinks to the document."""
    link_paragraph = document.add_paragraph()
    link_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # First Publications Link
    first_pub_run = link_paragraph.add_run('First Publications')
    first_pub_run.font.size = Pt(10)
    first_pub_run.underline = True
    first_pub_run.font.color.rgb = first_pub_run.font.color.rgb  # Default hyperlink color
    first_pub_run.hyperlink = '#first_publications'

    link_paragraph.add_run('\t\t')  # Adding 2 tab spaces

    # Granted Patents Link
    granted_pat_run = link_paragraph.add_run('Granted Patents')
    granted_pat_run.font.size = Pt(10)
    granted_pat_run.underline = True
    granted_pat_run.font.color.rgb = granted_pat_run.font.color.rgb  # Default hyperlink color
    granted_pat_run.hyperlink = '#granted_patents'

    document.add_paragraph()  # Line break


# ------------------------------
# Main Processing Functions
# ------------------------------

def create_patent_watch_doc(excel_path, output_path, template_path):
    """Generate the Patent Watch document from Excel data using a template."""
    document = Document(template_path)

    # Add Title
    title_paragraph = document.add_paragraph('2445_2446 - PATENT WATCH – (04-NOV-2024 to 15-NOV-2024)')
    title_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_run = title_paragraph.runs[0]
    title_run.bold = True
    title_run.font.size = Pt(14)
    title_run.font.all_caps = True

 

    # Add INDEX heading
    index_paragraph = document.add_paragraph('INDEX')
    index_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    index_run = index_paragraph.runs[0]
    index_run.bold = True
    index_run.font.size = Pt(12)
    index_run.all_caps = True
    index_run.underline = WD_UNDERLINE.SINGLE



    # Create Company Table
    first_row_data = ['CGG', 'WESTERNGECO', 'PGS', 'ION', 'BGP/CNPC/PETROCHINA']
    second_row_data = ['CHEVRON', 'HALLIBURTON/LANDMARK', 'FFA GEOTERIC', 'PARADIGM', 'WEATHERFORD', 'BAKER HUGHES']
    create_company_table(document, first_row_data, second_row_data)

    document.add_paragraph()  # Line break

    # Add Hyperlinks for First Publications & Granted Patents
    add_patent_section_links(document)

    # ------------------------------
    # Load Category Data from Excel
    # ------------------------------

    df_fp = pd.read_excel(excel_path, sheet_name='FP')
    df_fp.columns = df_fp.columns.str.strip()
    df_fp['Category'] = df_fp['Category'].astype(str).str.strip()

    df_grant = pd.read_excel(excel_path, sheet_name='Grant')
    df_grant.columns = df_grant.columns.str.strip()
    df_grant['Category'] = df_grant['Category'].astype(str).str.strip()

    categories = ['Seafloor', 'Land', 'Marine', 'Microseismic & Multiphysics',
                  'Processing', 'Reservoir', 'Geology', 'Data Management & Computing']

    # Split into two groups
    categories_group1, categories_group2 = categories[:5], categories[5:]

    # Process first five categories
    for category in categories_group1:
        process_category(document, category, df_fp, df_grant)

    document.add_paragraph()  # Line break

    # Create Company Table Again After First Five Categories
    create_company_table(document, first_row_data, second_row_data)

    document.add_paragraph()  # Line break

    # Process remaining three categories
    for category in categories_group2:
        process_category(document, category, df_fp, df_grant)

    # Save Document
    document.save(output_path)


# ------------------------------
# Process Category Function
# ------------------------------

def process_category(document, category, df_fp, df_grant):
    """Process a single category and insert related values into the document."""
    bullet_paragraph = document.add_paragraph()
    bullet_paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
    bullet_run = bullet_paragraph.add_run(u'\u2192 ' + category)
    bullet_run.font.size = Pt(10)
    bullet_run.bold = True

    fp_values = df_fp[df_fp['Category'].str.lower().str.contains(category.lower(), na=False)]['Publication No'].tolist()
    grant_values = df_grant[df_grant['Category'].str.lower().str.contains(category.lower(), na=False)]['Patent No'].tolist()

    values = fp_values + grant_values

    if values:
        num_rows = math.ceil(len(values) / 4)
        value_table = document.add_table(rows=num_rows, cols=4)
        value_table.autofit = True

        for idx, value in enumerate(values):
            row_idx, col_idx = divmod(idx, 4)
            cell = value_table.rows[row_idx].cells[col_idx]
            cell.text = str(value)
            cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
            apply_font_size(cell)

        set_table_borders(value_table)
        set_cell_size_and_alignment(value_table, cell_height=0.28)

    else:
        no_value_paragraph = document.add_paragraph(f'No records found for {category}')
        no_value_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

    document.add_paragraph()  # Line break after each category


def apply_font_size(cell, font_size=10):
    """Ensure the font size is set correctly inside a table cell."""
    try:
        cell.paragraphs[0].runs[0].font.size = Pt(font_size)
    except IndexError:
        run = cell.paragraphs[0].add_run()
        run.font.size = Pt(font_size)


# ------------------------------
# Script Execution
# ------------------------------

if __name__ == "__main__":
    excel_path = sys.argv[1]
    output_file = sys.argv[2]
    template_file = sys.argv[3]

    create_patent_watch_doc(excel_path, output_file, template_file)