import pandas as pd
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_UNDERLINE
from docx.enum.table import WD_ALIGN_VERTICAL
import sys
import math
import os
import logging
import requests
from io import BytesIO
from PIL import Image


# Set up logging
logging.basicConfig(level=logging.INFO, format='%(asctime)s - %(levelname)s - %(message)s')
logger = logging.getLogger(__name__)

def add_hyperlink(paragraph, text, url):
    """Add a hyperlink to a paragraph."""
    part = paragraph.part
    r_id = part.relate_to(url, 'http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink', is_external=True)
    
    hyperlink = OxmlElement('w:hyperlink')
    hyperlink.set(qn('r:id'), r_id)
    
    run = OxmlElement('w:r')
    rPr = OxmlElement('w:rPr')
    
    # Add underline
    underline = OxmlElement('w:u')
    underline.set(qn('w:val'), 'single')
    rPr.append(underline)
    
    # Add blue color
    color = OxmlElement('w:color')
    color.set(qn('w:val'), '0000FF')
    rPr.append(color)
    
    run.append(rPr)
    t = OxmlElement('w:t')
    t.text = text
    run.append(t)
    hyperlink.append(run)
    paragraph._element.append(hyperlink)

def set_column_widths(table, left_width, right_width):
    """Set the widths of the columns in a table."""
    for row in table.rows:
        row.cells[0].width = left_width
        row.cells[1].width = right_width

def add_table_borders(table):
    """Add borders to a table."""
    tbl = table._tbl  # Get the table XML element
    for cell in tbl.iter(qn('w:tc')):
        tcPr = cell.find(qn('w:tcPr'))
        if tcPr is None:
            tcPr = OxmlElement('w:tcPr')
            cell.insert(0, tcPr)
        tcBorders = OxmlElement('w:tcBorders')
        for border_name in ['top', 'left', 'bottom', 'right']:
            border = OxmlElement(f'w:{border_name}')
            border.set(qn('w:val'), 'single')
            border.set(qn('w:sz'), '4')  # Border size
            border.set(qn('w:space'), '0')
            border.set(qn('w:color'), '000000')  # Black color
            tcBorders.append(border)
        tcPr.append(tcBorders)

def create_first_publications_section(document, df, df_images):
    """Create the First Publications section in the document."""
    # Add heading
    heading_para = document.add_paragraph()
    heading_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    heading_run = heading_para.add_run("FIRST PUBLICATIONS")
    heading_run.bold = True
    heading_run.underline = True
    heading_run.font.size = Pt(11)
    
    # Add index only once at the beginning
    index_para = document.add_paragraph()
    index_para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    index_run = index_para.add_run("<<INDEX")
    index_run.font.size = Pt(10)
    
    # Define column widths
    left_col_width = Inches(1.38)
    right_col_width = Inches(5.61)
    
    # Iterate over each record in the DataFrame
    first_record = True
    for _, row in df.iterrows():
        # Only add page break after the first record
        if not first_record:
            document.add_page_break()
            
            # Add index at the top of each new page
            index_para = document.add_paragraph()
            index_para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            index_run = index_para.add_run("<<INDEX")
            index_run.font.size = Pt(10)
        
        first_record = False
        
        # Create a table with 2 columns
        table = document.add_table(rows=0, cols=2)
        table.style = 'Table Grid'  # Simple way to add borders
        
        # Set column widths directly
        table.columns[0].width = Inches(1.38)
        table.columns[1].width = Inches(5.61)
        
        # Add rows for each field
        fields = [
            ("Serial No", row['Serial No']),
            ("Publication No", row['Publication No']),
            ("Kind Code", row['Kind Code']),
            ("Title", row['Title']),
            ("Publication Date", row['Publication Date']),
            ("Earliest Priority Date", row['Earliest Priority Date']),
            ("Assignee", row['Assignee']),
            ("Inventors", row['Inventors']),
            ("Category", row['Category']),
            ("IPC", row['IPC']),
            ("PDF Document", row['Patent Link']),
            ("Abstract", row['Abstract'])
        ]
        
        for field_name, field_value in fields:
            row_cells = table.add_row().cells
            row_cells[0].text = field_name
            row_cells[0].vertical_alignment = WD_ALIGN_VERTICAL.TOP
            row_cells[1].text = str(field_value) if pd.notna(field_value) else ""
            row_cells[1].vertical_alignment = WD_ALIGN_VERTICAL.TOP
            
            # Set fixed height for all fields except Abstract and Image
            if field_name not in ["Abstract", "Image"]:
                row_cells[0].height = Inches(0.23)
                row_cells[1].height = Inches(0.23)
        
        # Add hyperlink for PDF Document
        if pd.notna(row['Patent Link']):
            pdf_para = table.rows[10].cells[1].paragraphs[0]  # PDF Document row
            pdf_para.clear()
            add_hyperlink(pdf_para, "Link", row['Patent Link'])
        
        # Lookup Image URL from Sheet1 using Family number
        image_url = None
        if "Family number" in row and pd.notna(row["Family number"]):
            image_row = df_images[df_images['Family number'] == row['Family number']]
            if not image_row.empty:
                image_url = image_row.iloc[0]['Image']

        # Insert Image Row Below Abstract
        if image_url:
            image_title_cell = table.add_row().cells
            image_title_cell[0].text = "Image"  # Left cell title
            download_and_insert_image(image_title_cell[1], image_url)  # Insert image in right cell

def download_and_insert_image(cell, image_url):
    """Downloads an image from a URL and inserts it into a cell."""
    try:
        response = requests.get(image_url, stream=True)
        if response.status_code == 200:
            img = Image.open(BytesIO(response.content))
            temp_img = BytesIO()
            img.save(temp_img, format='PNG')
            temp_img.seek(0)
            img_para = cell.add_paragraph()
            img_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = img_para.add_run()
            run.add_picture(temp_img, width=Inches(2))
    except Exception as e:
        print(f"Error downloading image: {e}")

def create_granted_patents_section(document, df_granted, df_images):
    """
    Creates the Granted Patents section.
    Expected columns in df_granted (Granted sheet):
      A - Serial No, B - Family number, C - Patent No, D - Kind Code, 
      E - Title, F - Publication Date, G - Earliest Priority, H - Assignee,
      I - Inventors, J - Category, K - IPC, L - URL (no title), M - File format,
      N - Patent Link, O - Abstract
    """
    # Insert a page break to start Granted Patents on a new page.
    document.add_page_break()
    
    # Add a heading for Granted Patents.
    heading_para = document.add_paragraph()
    heading_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    heading_run = heading_para.add_run("GRANTED PATENTS")
    heading_run.bold = True
    heading_run.underline = True
    heading_run.font.size = Pt(11)
    
    # Add index only once at the beginning
    index_para = document.add_paragraph()
    index_para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    index_run = index_para.add_run("<<INDEX")
    index_run.font.size = Pt(10)
    
    # Define Granted Patents headers
    granted_headers = [
        "Serial No", "Family number", "Patent No", "Kind Code",
        "Title", "Publication Date", "Earliest Priority", "Assignee",
        "Inventors", "Category", "IPC", "Patent Link", "Abstract", "Image"
    ]
    
    # Process each granted patent record.
    for idx, row in df_granted.iterrows():
        # Create a table for each record with one row per field.
        process_record(document, row.to_dict(), granted_headers, df_images, mode="granted")
        # Optionally add a spacing paragraph between records.
        document.add_paragraph("")

def process_record(document, record, headers, df_images, mode="first"):
    """
    Processes a single record and adds it to the document.
    
    :param document: The Word document object.
    :param record: A dictionary of the record data.
    :param headers: A list of headers for the table.
    :param df_images: DataFrame containing image URLs.
    :param mode: Mode of processing, either "first" or "granted".
    """
    # Create a table with 2 columns
    table = document.add_table(rows=0, cols=2)
    table.style = 'Table Grid'  # Simple way to add borders
    
    # Set column widths
    table.columns[0].width = Inches(1.38)
    table.columns[1].width = Inches(5.61)
    
    # Add rows for each field
    for header in headers:
        row_cells = table.add_row().cells
        row_cells[0].text = header
        row_cells[0].vertical_alignment = WD_ALIGN_VERTICAL.TOP
        row_cells[1].text = str(record.get(header, "")) if pd.notna(record.get(header, "")) else ""
        row_cells[1].vertical_alignment = WD_ALIGN_VERTICAL.TOP
        
        # Set fixed height for all fields except Abstract and Image
        if header not in ["Abstract", "Image"]:
            row_cells[0].height = Inches(0.23)
            row_cells[1].height = Inches(0.23)
    
    # Add hyperlink for PDF Document
    if pd.notna(record.get('Patent Link')):
        pdf_para = table.rows[headers.index('Patent Link')].cells[1].paragraphs[0]  # PDF Document row
        pdf_para.clear()
        add_hyperlink(pdf_para, "Link", record['Patent Link'])
    
    # Lookup Image URL from Sheet1 using Family number
    image_url = None
    if "Family number" in record and pd.notna(record["Family number"]):
        image_row = df_images[df_images['Family number'] == record['Family number']]
        if not image_row.empty:
            image_url = image_row.iloc[0]['Image']

    # Insert Image Row Below Abstract
    if image_url:
        image_title_cell = table.add_row().cells
        image_title_cell[0].text = "Image"  # Left cell title
        download_and_insert_image(image_title_cell[1], image_url)  # Insert image in right cell

def main():
    try:
        # Load Excel sheets
        excel_path = sys.argv[1] if len(sys.argv) > 1 else r'C:\Users\Ayman\Documents\Abhijit_mail_attachments\Test_PW.xlsm'
        df_fp = pd.read_excel(excel_path, sheet_name="First Publication")
        df_granted = pd.read_excel(excel_path, sheet_name="Granted")
        df_images = pd.read_excel(excel_path, sheet_name="Sheet1")
        
        # Load the template document
        template_file = "basic_page_template.docx"
        document = Document(template_file)
        
        # Create the First Publications section
        create_first_publications_section(document, df_fp, df_images)
        
        # Create the Granted Patents section
        create_granted_patents_section(document, df_granted, df_images)
        
        # Save the final document
        output_path = "part_4.docx"
        document.save(output_path)
        logger.info(f"Document saved: {output_path}")
        print(f"Successfully created document: {output_path}")
    
    except Exception as e:
        logger.error(f"Error generating document: {e}")
        print(f"Critical error: {e}")

if __name__ == "__main__":
    main()