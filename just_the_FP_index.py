import pandas as pd
from docx import Document
from docx.shared import Inches, Pt
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ROW_HEIGHT_RULE
import sys
from docx.shared import RGBColor

# Function to set table borders
def set_table_borders(table):
    tbl = table._tbl
    tblPr = tbl.tblPr
    borders = OxmlElement('w:tblBorders')
    for border_name in ['top', 'left', 'bottom', 'right', 'insideH', 'insideV']:
        border = OxmlElement(f'w:{border_name}')
        border.set(qn('w:val'), 'single')
        border.set(qn('w:sz'), '4')  # Size of the border
        border.set(qn('w:space'), '0')
        border.set(qn('w:color'), '000000')  # Black color for the border
        borders.append(border)
    tblPr.append(borders)


def add_hyperlink(paragraph, text, bookmark_name):
    """Add an internal hyperlink (bookmark link) in a Word document."""
    run = paragraph.add_run(text)
    run.font.color.rgb = RGBColor(0, 0, 255)  # Blue color for hyperlink
    run.underline = True

    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("w:anchor"), bookmark_name)  # Reference the bookmark name
    hyperlink.append(run._r)
    paragraph._element.append(hyperlink)


def create_first_publications_doc(excel_path, output_path, template_path):
    # Read Excel data from the 'First Publication' worksheet
    df = pd.read_excel(excel_path, sheet_name='First Publication')
    
    # Clean column names and ensure there are no extra spaces
    df.columns = df.columns.str.strip()
    
    # Clean the 'Category' column values
    df['Category'] = df['Category'].astype(str).str.strip()

    # Load the template document
    document = Document(template_path)

    # Define categories and their corresponding widths as two lists
    categories_list = ['Seafloor', 'Land', 'Marine', 'Microseismic & Multiphysics',
                    'Processing', 'Reservoir', 'Geology', 'Data Management & Computing',
                    'Downhole']
    widths_list = [0.69, 0.58, 0.66, 0.93, 0.78, 0.73, 0.69, 0.95, 0.76]

    # Create index table with placeholders
    index_table = document.add_table(rows=1, cols=len(categories_list))

    # Disable auto-fit by setting a fixed layout for the table
    tbl = index_table._tbl
    tblPr = tbl.tblPr
    tblLayout = OxmlElement('w:tblLayout')
    tblLayout.set(qn('w:type'), 'fixed')
    tblPr.append(tblLayout)

    index_row = index_table.rows[0].cells
    for i, (category, width) in enumerate(zip(categories_list, widths_list)):
        cell = index_row[i]
        paragraph = cell.paragraphs[0]
        run = paragraph.add_run(category)
        run.font.size = Pt(10)
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # Set column width
        cell.width = Inches(width)
        tcPr = cell._tc.get_or_add_tcPr()
        tcW = OxmlElement("w:tcW")
        tcW.set(qn("w:w"), str(int(width * 1440)))  # Convert inches to twips
        tcW.set(qn("w:type"), "dxa")
        tcPr.append(tcW)



    
    # Set table borders for index table
    set_table_borders(index_table)

    document.add_paragraph()  # Line break

    # Add main heading
    heading_paragraph = document.add_paragraph('FIRST PUBLICATIONS')
    heading_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    heading_run = heading_paragraph.runs[0]
    heading_run.bold = True
    heading_run.font.size = Pt(12)

    # Create main data table with specified column widths
    table = document.add_table(rows=1, cols=5)
    table.autofit = False  # Disable autofit to manually set column widths

    # Set column widths
    column_widths = [Inches(0.45), Inches(1.13), Inches(2.43), Inches(1.35), Inches(1.4)]

    # Apply column widths to the table
    table.autofit = False  # Disable auto-fit to manually set column widths

    # Set column headers
    headers = ['Sl No', 'Publication No', 'Title', 'Assignee', 'Inventors']
    header_row = table.rows[0]

    for i, (header, width) in enumerate(zip(headers, column_widths)):
        cell = header_row.cells[i]
        cell.text = header

        paragraph = cell.paragraphs[0]
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER  # Center align headers

        run = paragraph.runs[0] if paragraph.runs else paragraph.add_run()
        run.bold = True
        run.font.size = Pt(10)

        # Apply width using cell._tc XML
        cell._tc.get_or_add_tcPr().append(OxmlElement("w:tcW"))
        cell._tc.get_or_add_tcPr().find(qn("w:tcW")).set(qn("w:w"), str(int(width.inches * 1440)))  # Convert inches to twips
        cell._tc.get_or_add_tcPr().find(qn("w:tcW")).set(qn("w:type"), "dxa")


        paragraph = cell.paragraphs[0]
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        paragraph.runs[0].bold = True
        paragraph.runs[0].font.size = Pt(10)
        paragraph.space_after = Pt(0)
        paragraph.space_before = Pt(0)

    # Process each category
    for category in categories_list:
        # Add category heading row
        cat_row = table.add_row()
        cat_row.height = Inches(0.24) 
        cat_row.height_rule = WD_ROW_HEIGHT_RULE.EXACTLY  # Ensures the row height is fixed

        cat_cell = cat_row.cells[0]
        cat_cell.merge(cat_row.cells[-1])
        cat_cell.text = f'> {category.upper()}'
        paragraph = cat_cell.paragraphs[0]
        paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
        paragraph.runs[0].bold = True
        paragraph.runs[0].font.size = Pt(10)

        # Filter the DataFrame for rows matching the current category (case-insensitive)
        cat_data = df[df['Category'].str.lower().str.contains(category.lower())]

        # Add rows for the data
        for _, row in cat_data.iterrows():
            data_row = table.add_row()
            data_row.cells[0].text = str(row.get('Serial No', ''))

            pub_no = str(row.get('Publication No', ''))
            cell = data_row.cells[1]  # Assuming Patent No is in the second column
            
            if pub_no:
                p = cell.paragraphs[0]
                add_hyperlink(p, pub_no, pub_no.strip())  # Use Publication No as the bookmark name
            else:
                cell.text = pub_no  # Just display the text if no number exists

            data_row.cells[2].text = str(row.get('Title', ''))
            data_row.cells[3].text = str(row.get('Assignee', ''))
            data_row.cells[4].text = str(row.get('Inventors', ''))

            # Set column widths and font size for data rows
            for i, cell in enumerate(data_row.cells):
                cell.width = column_widths[i]
                for paragraph in cell.paragraphs:
                    paragraph.space_after = Pt(0)
                    paragraph.space_before = Pt(0)
                    for run in paragraph.runs:
                        run.font.size = Pt(10)

    # Set table borders for main data table
    set_table_borders(table)

    # Save document
    document.save(output_path)

# Usage example:
# create_first_publications_doc(
#     'C:/Users/Ayman/Documents/Abhijit_mail_attachments/Test_PW.xlsm',
#     'part_4.docx',
#     'basic_page_template.docx'
# )



if __name__ == "__main__":
    excel_path = sys.argv[1]
    output_file = sys.argv[2]
    template_file = sys.argv[3]

    create_first_publications_doc(excel_path, output_file, template_file)

# # Run the function directly to generate output
# excel_path = r"C:\Users\Ayman\Documents\Abhijit_mail_attachments\Test_PW.xlsm"
# output_file = "FP_output.docx"
# template_file = "basic_page_template.docx"

# create_first_publications_doc(excel_path, output_file, template_file)

print(f"First Publications index has been generated: {output_file}")
