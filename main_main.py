import subprocess
import shutil
from docx import Document
from docxcompose.composer import Composer
from docx.oxml.ns import qn

# Paths
excel_path = "C:/Users/Ayman/Documents/Abhijit_mail_attachments/Test_PW.xlsm"
output_file = "final_patent_watch.docx"
template_file = "basic_page_template.docx"

# List of script and output document pairs
scripts = [
    ("the_first_2_pages.py", "part_1.docx"),
    ("just_the_FP_index.py", "part_2.docx"),
    ("just_the_GP_index.py", "part_3.docx"),
    ("first_publications_pages_generator.py", "part_4.docx")
]

# Run each script to generate its document
for script, part_file in scripts:
    try:
        print(f"Running {script} to generate {part_file}...")
        subprocess.run(["python", script, excel_path, part_file, template_file], check=True)
    except subprocess.CalledProcessError as e:
        print(f"Error running {script}: {e}")

# # Ensure we use the correctly formatted output from `so_we_cry.py`
# print("Running so_we_cry.py to generate part4.docx...")
# try:
#     subprocess.run(["python", "so_we_cry.py", excel_path], check=True)
# except subprocess.CalledProcessError as e:
#     print(f"Error running so_we_cry.py: {e}")

# Function to check if the document ends at the top of a new page
def is_cursor_at_top_of_page(doc):
    if not doc.paragraphs:
        return True  # Empty document = start of new page

    last_para = doc.paragraphs[-1]
    for run in last_para.runs:
        for child in run._element:
            if child.tag == qn('w:br') and child.get(qn('w:type')) == "page":
                return True  # Page break found

    return False  # No page break found, so not at the top of a page.

# Merge all generated documents while preserving bookmarks and hyperlinks
def merge_documents(output_file, parts):
    master = Document(parts[0])  # Start with the first document
    composer = Composer(master)

    for part in parts[1:]:
        # Remove trailing blank pages before adding a new section
        while master.paragraphs and not master.paragraphs[-1].text.strip():
            p = master.paragraphs[-1]
            master._element.body.remove(p._element)  # ✅ Delete empty paragraphs

    if not is_cursor_at_top_of_page(master):
        master.add_page_break()

    doc_to_append = Document(part)

    # Instead of using composer.append(), manually append elements to preserve hyperlinks
    for element in doc_to_append.element.body:
        master.element.body.append(element)  # ✅ Preserves bookmarks and hyperlinks

    # Reapply table style if lost
    for table in master.tables:
        table.style = 'Table Grid'

    master.save(output_file)
    print(f"Final document '{output_file}' created successfully!")

# Execute the merge
merge_documents(output_file, [
    "part_1.docx",  # ✅ Title pages
    "part_2.docx",  # ✅ FP Index 
    "part_3.docx",  # ✅ GP Index 
    "part_4.docx",   # ✅ First Publications & Granted Patents
])

# import subprocess
# from docx import Document
# from docxcompose.composer import Composer
# from docx.oxml.ns import qn
# from docx.shared import Pt
# from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK

# # Paths
# excel_path = "C:/Users/Ayman/Documents/Abhijit_mail_attachments/Test_PW.xlsm"
# output_file = "final_patent_watch.docx"
# template_file = "basic_page_template.docx"

# # List of script and output document pairs
# scripts = [
#     ("the_first_2_pages.py", "part_1.docx"),
#     ("just_the_FP_index.py", "part_2.docx"),
#     ("just_the_GP_index.py", "part_3.docx"),
#     ("first_publications_pages_generator.py", "part_4.docx"),
#     ("granted_patents_pages_generator.py", "part_5.docx"),
# ]

# # Run each script to generate its document
# for script, part_file in scripts:
#     try:
#         print(f"Running {script} to generate {part_file}...")
#         subprocess.run(["python", script, excel_path, part_file, template_file], check=True)
#     except subprocess.CalledProcessError as e:
#         print(f"Error running {script}: {e}")

# def merge_documents(output_file, parts):
#     master = Document(parts[0])  # Start with the first document
    
#     for part in parts[1:]:
#         doc_to_append = Document(part)
        
#         # Add a page break before appending new content
#         if master.paragraphs:
#             last_paragraph = master.add_paragraph()
#             run = last_paragraph.add_run()
#             run.add_break(WD_BREAK.PAGE)
        
#         # Append each element from the document
#         for element in doc_to_append.element.body:
#             master.element.body.append(element)
    
#     # Clean up any consecutive page breaks
#     for i in range(len(master.paragraphs) - 1, 0, -1):
#         current_para = master.paragraphs[i]
#         prev_para = master.paragraphs[i-1]
#         if not current_para.text.strip() and not prev_para.text.strip():
#             p = master.paragraphs[i]
#             p._element.getparent().remove(p._element)
    
#     # Ensure table formatting is preserved
#     for table in master.tables:
#         table.style = 'Table Grid'
        
#     master.save(output_file)
#     print(f"Final document '{output_file}' created successfully!")

# # Execute the merge
# merge_documents(output_file, [
#     "part_1.docx",  # Title pages
#     "part_2.docx",  # FP Index
#     "part_3.docx",  # GP Index
#     "part_4.docx",  # First Publications
#     "part_5.docx"   # Granted Patents
# ])