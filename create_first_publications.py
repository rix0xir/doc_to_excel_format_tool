import pandas as pd
from docx import Document
from docx.shared import Inches, RGBColor, Pt
from docx.oxml.ns import qn, nsdecls
from docx.oxml import OxmlElement, parse_xml
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_ROW_HEIGHT_RULE
import requests
from io import BytesIO


# --------------------
# Helper functions
# --------------------

def resize_image(image_stream, fixed_height=4.5):
    """
    Resize the image to a fixed height while maintaining the aspect ratio.
    """
    img = Image.open(image_stream)
    original_width, original_height = img.size
    aspect_ratio = original_width / original_height
    new_height = int(fixed_height * 1440)  # Convert inches to twips
    new_width = int(new_height * aspect_ratio)
    resized_img = img.resize((new_width, new_height), Image.ANTIALIAS)
    return resized_img, new_width, new_height

def set_cell_margins(cell, top=20, start=20, bottom=20, end=20):
    """
    Set cell margins (padding) in twips.
    Default margins are reduced (20 twips ≈ 0.014") so that rows can shrink.
    """
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    # Remove any existing margins
    for m in ('w:topMar', 'w:startMar', 'w:bottomMar', 'w:endMar'):
        for node in tcPr.findall(m, namespaces={'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'}):
            tcPr.remove(node)
    cell_margins = parse_xml(
        r'<w:tcMar %s>'
        r'<w:top w:w="%s" w:type="dxa"/>' 
        r'<w:start w:w="%s" w:type="dxa"/>' 
        r'<w:bottom w:w="%s" w:type="dxa"/>' 
        r'<w:end w:w="%s" w:type="dxa"/>'
        r'</w:tcMar>' % (nsdecls('w'), top, start, bottom, end)
    )
    tcPr.append(cell_margins)

def set_cell_width(cell, width_twips):
    """
    Set a cell’s width explicitly in twips.
    """
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    for child in tcPr.findall(qn('w:tcW')):
        tcPr.remove(child)
    tcW = OxmlElement('w:tcW')
    tcW.set(qn('w:w'), str(width_twips))
    tcW.set(qn('w:type'), 'dxa')
    tcPr.append(tcW)

def set_paragraph_font(paragraph, font_name='Calibri', font_size=Pt(10)):
    """
    Set the font for all runs in a paragraph.
    """
    for run in paragraph.runs:
        run.font.name = font_name
        run.font.size = font_size

# --------------------
# Define table dimensions:
# --------------------
# Convert inches to twips (1 inch = 1440 twips)
left_width_twips = int(1.43 * 1440)  # ≈2059 twips
right_width_twips = int(5.06 * 1440)  # ≈7286 twips
total_width_twips = left_width_twips + right_width_twips  # ≈9345 twips

# --------------------
# Load Excel Data
# --------------------
excel_path = 'C:/Users/Ayman/Documents/Abhijit_mail_attachments/Test_PW.xlsm'
first_pub_df = pd.read_excel(excel_path, sheet_name='First Publication')
sheet1_df = pd.read_excel(excel_path, sheet_name='Sheet1')

# --------------------
# Open Document and add Title/Index
# --------------------
doc_path = 'basic_page_template.docx'
doc = Document(doc_path)

# Title (set to 11 pt)
title = doc.add_paragraph()
title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
title_run = title.add_run('FIRST PUBLICATIONS')
title_run.bold = True
title_run.underline = True
title_run.font.size = Pt(11)

# Index link
index_paragraph = doc.add_paragraph()
index_paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
index_run = index_paragraph.add_run('<< INDEX')
index_run.font.color.rgb = RGBColor(0, 0, 255)
index_run.font.underline = True
hyperlink = OxmlElement('w:hyperlink')
hyperlink.set(qn('r:id'), 'rId1')
hyperlink.append(index_run._r)
index_paragraph._element.append(hyperlink)

# --------------------
# Process Each Record and Create Table
# --------------------
headings = [
    'Serial No', 'Publication No', 'Kind Code', 'Title', 'Publication Date',
    'Earliest Priority Date', 'Assignee', 'Inventors', 'Category', 'IPC',
    'Patent Link', 'Abstract', 'Image'
]

# Namespace mapping for XML operations
nsmap = {'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'}

for _, row in first_pub_df.iterrows():
    row_data = row.to_dict()  # Convert row to dictionary

    # Add index link on each new page (from second page onward)
    if doc.paragraphs[-1].text != '<< INDEX':
        index_paragraph = doc.add_paragraph()
        index_paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
        index_run = index_paragraph.add_run('<< INDEX')
        index_run.font.color.rgb = RGBColor(0, 0, 255)
        index_run.font.underline = True
        hyperlink = OxmlElement('w:hyperlink')
        hyperlink.set(qn('r:id'), 'rId1')
        hyperlink.append(index_run._r)
        index_paragraph._element.append(hyperlink)

    # Create table with rows = number of headings and 2 columns
    table = doc.add_table(rows=len(headings), cols=2)
    table.style = 'Table Grid'

    # --- Set table properties: fixed layout and grid columns ---
    tbl = table._tbl
    tblPr = tbl.find('./w:tblPr', namespaces=nsmap)
    if tblPr is None:
        tblPr = OxmlElement('w:tblPr')
        tbl.insert(0, tblPr)
    tblW = OxmlElement('w:tblW')
    tblW.set(qn('w:w'), str(total_width_twips))
    tblW.set(qn('w:type'), 'dxa')
    tblPr.append(tblW)
    tblLayout = OxmlElement('w:tblLayout')
    tblLayout.set(qn('w:type'), 'fixed')
    tblPr.append(tblLayout)
    tblGrid = OxmlElement('w:tblGrid')
    gridCol1 = OxmlElement('w:gridCol')
    gridCol1.set(qn('w:w'), str(left_width_twips))
    tblGrid.append(gridCol1)
    gridCol2 = OxmlElement('w:gridCol')
    gridCol2.set(qn('w:w'), str(right_width_twips))
    tblGrid.append(gridCol2)
    tbl.append(tblGrid)

    # --- Fill non-image rows (rows 0 to 11; Abstract is row 11) ---
    values = [
        str(row_data.get('Serial No', '')),
        str(row_data.get('Publication No', '')),
        str(row_data.get('Kind Code', '')),
        str(row_data.get('Title', '')),
        str(row_data.get('Publication Date', '')),
        str(row_data.get('Earliest Priority Date', '')),
        str(row_data.get('Assignee', '')),
        str(row_data.get('Inventors', '')),
        str(row_data.get('Category', '')),
        str(row_data.get('IPC', '')),
        str(row_data.get('Patent Link', '')),
        str(row_data.get('Abstract', ''))
    ]
    # For rows 0 to 10 (all except Abstract and Image)
    for i in range(len(headings) - 1):
        cell_left = table.cell(i, 0)
        cell_right = table.cell(i, 1)
        cell_left.text = headings[i]
        cell_right.text = values[i]
        cell_left.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        cell_right.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        # Force explicit cell widths
        set_cell_width(cell_left, left_width_twips)
        set_cell_width(cell_right, right_width_twips)
        # For all rows except the Abstract (row 11) and Image (row 12), reduce cell margins
        set_cell_margins(cell_left, top=20, start=20, bottom=20, end=20)
        set_cell_margins(cell_right, top=20, start=20, bottom=20, end=20)
        # Set font for cell paragraphs to Calibri 10
        for paragraph in cell_left.paragraphs:
            set_paragraph_font(paragraph, font_name='Calibri', font_size=Pt(10))
        for paragraph in cell_right.paragraphs:
            set_paragraph_font(paragraph, font_name='Calibri', font_size=Pt(10))
        # For rows 0 to 10 (excluding Abstract at index 11), set fixed height of 0.28"
        if i != 11:  # row 11 is Abstract; leave it auto-sized
            table.rows[i].height = Inches(0.28)
            table.rows[i].height_rule = WD_ROW_HEIGHT_RULE.EXACTLY

    # --- Last row for the image (row index 12) ---
    image_title_cell = table.cell(len(headings) - 1, 0)
    image_title_cell.text = 'Image'
    image_title_cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    set_cell_width(image_title_cell, left_width_twips)
    set_cell_margins(image_title_cell, top=20, start=20, bottom=20, end=20)
    for paragraph in image_title_cell.paragraphs:
        set_paragraph_font(paragraph, font_name='Calibri', font_size=Pt(10))
    
    image_cell = table.cell(len(headings) - 1, 1)
    image_cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    set_cell_width(image_cell, right_width_twips)
    set_cell_margins(image_cell, top=20, start=20, bottom=20, end=20)
    for paragraph in image_cell.paragraphs:
        set_paragraph_font(paragraph, font_name='Calibri', font_size=Pt(10))
    
    # --- Insert image into right cell of the last row ---
    family_number = row_data.get('Family number', '')
    if family_number and not sheet1_df[sheet1_df['Family number'] == family_number].empty:
        image_link = sheet1_df.loc[sheet1_df['Family number'] == family_number, 'Image'].values[0]
        try:
            response = requests.get(image_link)
            response.raise_for_status()
            image_stream = BytesIO(response.content)
            p = image_cell.paragraphs[0]
            p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
            run = p.add_run()
            run.add_picture(image_stream, width=Inches(2))
        except Exception as e:
            print(f"Error fetching image for Family number {family_number}: {e}")

    # --- Set row heights for remaining rows ---
    # Rows 0 to 10 are already set to 0.28" (except row 11 - Abstract and row 12 - Image)
    # For Abstract row (index 11), do not set a fixed height (allow auto height)
    # The Image row (index 12) is left unchanged.
    
    # Add a page break after each record
    doc.add_page_break()

# Save the document
doc.save('updated_publications.docx')
print("Document successfully created: 'updated_publications.docx'")
